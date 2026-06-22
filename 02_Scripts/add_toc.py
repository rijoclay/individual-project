# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

TEMPLATE = r"C:\Users\richa\Documents\Sem.6\DatEng_Individual_Project\Materials\PROPOSAL_individual project.docx"
REPORT_DIR = r"C:\Users\richa\Documents\Sem.6\DatEng_Individual_Project\individual_project_2\07_Reports"

def add_toc(doc, lang):
    """Insert a TOC field on a fresh page after cover."""
    doc.add_page_break()

    title = "Table of Contents" if lang == "EN" else "Daftar Isi"
    placeholder = (
        "[Right-click and select 'Update Field' to populate Table of Contents]"
        if lang == "EN"
        else "[Klik kanan dan pilih 'Update Field' untuk mengisi Daftar Isi]"
    )

    # Heading
    p = doc.add_paragraph()
    p.alignment = 1  # CENTER
    run = p.add_run(title)
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # TOC field
    p2 = doc.add_paragraph()

    # begin
    run_b = p2.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_b._r.append(fld_begin)

    # instrText
    run_i = p2.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    run_i._r.append(instr)

    # separate
    run_s = p2.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_s._r.append(fld_sep)

    # placeholder text
    run_t = p2.add_run(placeholder)
    run_t.font.name = "Times New Roman"
    run_t.font.size = Pt(12)
    run_t.font.color.rgb = RGBColor(0, 0, 0)

    # end
    run_e = p2.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_e._r.append(fld_end)

def make_proposal(lang):
    doc = Document(TEMPLATE)
    add_toc(doc, lang)
    name = f"Richard_Clay_{lang}.docx"
    doc.save(os.path.join(REPORT_DIR, name))
    print(f"{lang} proposal with TOC saved to {name}")

if __name__ == "__main__":
    make_proposal("EN")
    make_proposal("ID")
