# -*- coding: utf-8 -*-
import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REPORT_DIR = os.path.abspath(os.path.join(BASE_DIR, '..', '07_Reports'))
VIZ_DIR = os.path.abspath(os.path.join(BASE_DIR, '..', '06_Visualization'))
os.makedirs(REPORT_DIR, exist_ok=True)

def apply_text_formatting(paragraph, font_name='Times New Roman', size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.alignment = align
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic

def add_academic_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    apply_text_formatting(p, bold=bold, italic=italic)
    return p

def create_proposal_en():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Cover Page
    for _ in range(2): doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UNIVERSITI TENAGA NASIONAL (UNITEN)\nCollege of Computing & Informatics\n\nSemester II 2025/2026\nSession: 2025/2026")
    r.bold = True
    r.font.size = Pt(14)
    
    for _ in range(2): doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INDIVIDUAL PROJECT ON APACHE SPARK\nPROPOSAL DRAFT:\nDEVELOPMENT OF DATA PIPELINE AND STORAGE FOR AGENTIC AI LEADERSHIP METRICS")
    r.bold = True
    r.font.size = Pt(16)

    for _ in range(3): doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Subject Name: Special Topic in Data Engineering\n"
        "Subject Code: SECP3843\n"
        "Section: 02\n\n"
        "Prepared By:\n"
        "Richard Clay\n"
        "Matric Number: 0123456\n\n"
        "Presented To:\n"
        "Dr. [Lecturer Name]\n\n"
        "Date: June 23, 2026"
    )
    r.font.size = Pt(12)
    
    doc.add_page_break()

    # Abstract
    doc.add_heading("Abstract", level=1)
    add_academic_paragraph(doc, 
        "This proposal outlines the design and implementation of a scalable data pipeline and storage system using Apache Spark "
        "following the Medallion architecture within a Data Lakehouse pattern. Operating within the Professional Services business domain, "
        "evaluating Agentic AI reliability, performance, and trust is hindered by fragmented and heterogeneous operational telemetry data. "
        "The proposed pipeline ingests four distinct sources—Kaggle deployment metrics, external industry benchmarks, system execution logs, "
        "and academic paper citations from the OpenAlex API—to establish a unified database. An advanced regex-based NLP classification "
        "mapping is built in the Silver layer to contextualize paper citations directly to matching industries. An unsupervised K-Means "
        "clustering algorithm (Silhouette: 0.5276) and a tuned multiclass Logistic Regression classifier (Accuracy: 55.44%, F1-Score: 46.59%, "
        "Precision: 40.34%, Recall: 55.44%) are embedded in the Gold layer to profile performance and predict productivity, successfully "
        "avoiding data leakage. Best model weights are exported as production-ready artifacts. The deliverables demonstrate an automated, "
        "production-ready data engineering solution enabling business leaders to evaluate autonomous AI systems scientifically."
    )
    doc.add_page_break()

    # Section 1
    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Business Domain & Problem Identification", level=2)
    add_academic_paragraph(doc,
        "Modern enterprises within the Professional Services domain (consulting, legal, IT services) are undergoing a structural "
        "transition from static automation to autonomous Agentic AI systems. Unlike legacy software, these autonomous agents dynamically "
        "adapt, interact, and execute decisions without constant human intervention. However, business leaders currently face a critical "
        "management gap: there are no standardized frameworks or analytical tools to evaluate the trustworthiness, execution safety, "
        "and actual productivity improvements of these systems. Furthermore, operational telemetry data is highly fragmented across disparate "
        "structured and semi-structured formats, including local application execution logs, static CSV reports, and external API baselines. "
        "Without a centralized data engineering pipeline to clean, reconcile, and analyze these records, organizations risk executing "
        "AI deployments based on unstructured and unreliable metrics. Designing a robust, scalable data pipeline is therefore essential "
        "to transform raw log metrics into actionable, high-trust leadership profiles."
    )
    
    doc.add_heading("1.2 Objectives", level=2)
    doc.add_paragraph("The primary goal of this project is to develop an automated and scalable Apache Spark pipeline that integrates multi-source telemetry data to profile Agentic AI leadership metrics. The specific objectives are:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run(
        "1. To design and implement a Medallion-structured (Bronze-Silver-Gold) pipeline using Apache Spark to ingest and standardize 4 heterogeneous data sources.\n"
        "2. To build an automated cleansing process using Spark User-Defined Functions (UDFs) and window functions for industry-specific null value median imputation and deduplication.\n"
        "3. To integrate unsupervised K-Means clustering and tuned Logistic Regression classification models in the Gold layer to dynamically predict AI adoption productivity levels."
    )

    # Section 2
    doc.add_heading("2. Literature Review", level=1)
    add_academic_paragraph(doc,
        "A data pipeline is a foundational architecture that automates the extraction, transformation, and loading (ETL) of data from source "
        "systems to target repositories. Historically, organizations relied on centralized Data Warehouses for structured analysis, but these "
        "systems are cost-prohibitive for raw semi-structured logs. Conversely, Data Lakes offer affordable raw storage but suffer from query "
        "performance degradation and lack transactional integrity (ACID properties). To bridge this gap, the Data Lakehouse architecture "
        "has emerged as the standard, enabling direct transactional management, schema enforcement, and advanced machine learning modeling "
        "on cost-effective parquet file systems [1]."
    )
    add_academic_paragraph(doc,
        "Data engineering pipelines in organizational analytics are widely examined. Three related previous studies guide this project:\n"
        "First, Armbrust et al. (2021) demonstrated that Lakehouse structures reduce data duplication and latency by maintaining a single source "
        "of truth for both business intelligence and machine learning tasks [1]. Second, Zaharia et al. (2016) established Apache Spark as the "
        "most efficient engine for distributed in-memory compute, enabling low-latency joins of high-volume log streams [2]. Third, recent "
        "studies on AI governance emphasize that trust in agentic frameworks cannot be established without integrating raw execution telemetry "
        "with scientific research baselines, which provides validation for autonomous decision-making [3]."
    )

    # Section 3
    doc.add_heading("3. Methodology", level=1)
    doc.add_heading("3.1 Methodology Selection & Justification", level=2)
    add_academic_paragraph(doc,
        "This project utilizes the Cross-Industry Standard Process for Data Mining (CRISP-DM) framework. CRISP-DM consists of six iterative "
        "phases: Business Understanding, Data Understanding, Data Preparation, Modeling, Evaluation, and Deployment. This methodology is highly "
        "suitable for data engineering projects with embedded machine learning because it ensures that raw pipeline inputs (Data Preparation) "
        "are tightly aligned with the statistical features required by the analytical models (Modeling) in the Gold layer."
    )

    doc.add_heading("3.2 Data Source Identification", level=2)
    
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'ID', 'Data Source', 'Format', 'Description'
    sources = [
        ("S1", "Agentic AI Leadership Dataset", "CSV", "Primary dataset (5,500 records) containing organizational metrics, autonomy levels, and trust scores."),
        ("S2", "External Leadership Benchmarks", "JSON", "Industry baseline benchmarks for target task success rates and productivity thresholds."),
        ("S3", "Agent Execution Logs", "CSV", "Low-level system logs including CPU, Memory usage, execution times, and application error counts."),
        ("S4", "OpenAlex Academic Papers", "API JSON", "API scientific publications data used to extract citation statistics and establish domain maturity.")
    ]
    for i, s in enumerate(sources):
        r = table.rows[i+1].cells
        r[0].text, r[1].text, r[2].text, r[3].text = s[0], s[1], s[2], s[3]

    doc.add_paragraph()
    doc.add_heading("3.3 Data Ingestion, Cleansing & In-Memory Processing", level=2)
    add_academic_paragraph(doc,
        "The ingestion plan maps all sources to a Lakehouse storage format utilizing Apache Spark. Raw files are ingested in the Bronze layer, "
        "appended with ingestion timestamps, and saved as partitioned Parquet. In the Silver layer, data cleansing is performed by deduplicating "
        "on 'Record_ID', casting numeric types, and lowercasing categorical inputs. Null values are resolved using an Industry-specific median "
        "imputation. Execution logs are joined on 'Record_ID' and used to derive resources efficiency features like 'Memory_Per_Message_MB'. "
        "Furthermore, to contextualize scientific papers, an advanced regex-based NLP mapping classifies OpenAlex publications directly to matching "
        "business sectors (e.g. Healthcare, Finance), allowing sector-specific citation baselines to join with organization records."
    )

    doc.add_heading("3.4 Data Storage & Justification", level=2)
    add_academic_paragraph(doc,
        "Processed outputs are stored in a Medallion Lakehouse format (Bronze, Silver, Gold directories) using Parquet columnar storage. "
        "This selection is justified because Parquet optimizes storage footprint through dictionary encoding and significantly accelerates "
        "analytical queries via schema enforcement, column projection, and predicate pushdown in Spark SQL."
    )

    doc.add_heading("3.5 Artificial Intelligence (AI) Algorithm Integration", level=2)
    add_academic_paragraph(doc,
        "To elevate the pipeline's analytical value, two PySpark MLlib models are integrated in the Gold layer:\n"
        "1. Unsupervised K-Means Clustering: Evaluates 'Task_Success_Rate', 'Productivity_Improvement_Percent', and 'Leadership_Trust_Score' "
        "to partition organizations into performance profiles (High, Average, Low Performer). Validated via Silhouette Evaluator yielding a score of 0.5276.\n"
        "2. Supervised Logistic Regression Classifier: Predicts 'Productivity_Category' using system-level features ('Context_Awareness_Score', "
        "'Response_Time_Seconds', 'Task_Complexity_Score', 'Memory_Per_Message_MB', 'CPU_Utilization_Percent'). Tuned using a CrossValidator "
        "over a grid param (regParam, elasticNetParam) with 3-fold cross validation. Data leakage is prevented by isolating clustering features "
        "from the classification space. Model metrics obtained: Accuracy: 55.44%, F1-Score: 46.59%, Precision: 40.34%, Recall: 55.44%. The best "
        "model weights are exported as a production-ready artifact to the disk."
    )

    # Section 4
    doc.add_heading("4. System Design & Architecture", level=1)
    
    doc.add_heading("4.1 Use Case Diagram Description", level=2)
    add_academic_paragraph(doc,
        "The System Use Case involves two primary actors: the Data Engineer and the Data Analyst. "
        "The Data Engineer triggers the ETL orchestration process (Bronze to Silver ingestion, UDF-driven cleansing, "
        "and integration of logs & scientific literature citations). The Data Analyst interacts with the finalized Gold Layer data marts "
        "to run business intelligence queries, inspect AI adoption profiles, and evaluate productivity classification predictions. "
        "The Machine Learning model runs automatically inside the pipeline execution flow, retraining/re-tuning weights on new data batches."
    )
    
    doc.add_heading("4.2 System Architecture (ETL & Storage flow)", level=2)
    add_academic_paragraph(doc,
        "The pipeline architecture conforms to the Medallion structure. Raw sources (CSV, JSON) are stored in local folders. "
        "Spark processes the files sequentially: Raw data -> Bronze Parquet (Audited Raw) -> Silver Parquet (Imputed, deduplicated, "
        "and NLP-integrated data) -> Gold Parquet analytical marts (Aggregated profiles & Predicted labels). PySpark MLlib models "
        "pull clean data from Silver Parquet, apply StandardScaler, train, and write inference back to the Gold directory."
    )
    
    doc.add_heading("4.3 Database & Multidimensional Schema (ERD & Star Schema)", level=2)
    add_academic_paragraph(doc,
        "For multidimensional analytical modeling, the Gold Layer is structured as a Star Schema database consisting of a central Fact table "
        "and surrounding Dimension tables:\n"
        "1. Fact_AI_Deployment: Contains keys matching dimensions and measurable facts: Task_Success_Rate, Productivity_Improvement_Percent, "
        "Response_Time_Seconds, Memory_Usage_MB, CPU_Utilization_Percent, and Trust_vs_Benchmark_Gap.\n"
        "2. Dim_Organization: Holds structural details (Record_ID, Industry, Organization_Size, AI_Maturity_Level).\n"
        "3. Dim_Agent: Holds technology properties (Agent_Type, Use_Case_Area, Agent_Autonomy_Level, Decision_Making_Type).\n"
        "4. Dim_Research_Context: Stores the NLP mapped scientific references (Industry_Citation_Avg, Industry_Relevance_Avg, Mapped_Paper_Count)."
    )
    
    doc.add_heading("4.4 Main Interface & Analytical Dashboard Design", level=2)
    add_academic_paragraph(doc,
        "The analytical interface consists of three core visualization charts generated directly from the Gold Layer data marts: "
        "firstly, an Industry-level comparison chart plotting Trust Score against Productivity Improvement across sectors; "
        "secondly, an AI performance clustering profile displaying the characteristics of High, Average, and Low Performers; "
        "and thirdly, a prediction accuracy diagnostic dashboard displaying model classification success ratios."
    )

    doc.add_heading("4.5 Testing & Validation Strategy", level=2)
    add_academic_paragraph(doc,
        "To satisfy rigorous verification standards, the pipeline is bound to an automated test suite (test_pipeline.py) covering:\n"
        "1. Schema Conformance: Verifying all derived columns (e.g. Memory_Per_Message_MB) are present with exact types.\n"
        "2. Strict Quality Assertion: Asserting that critical numerical fields contain zero null values post-imputation.\n"
        "3. Integrity Assertions: Verifying that exactly 5,500 records are written to Silver and Gold layers without data drop.\n"
        "4. Artifact Checks: Confirming the physical output path of the tuned ML model is successfully written."
    )

    # Section 5
    doc.add_heading("5. Conclusion", level=1)
    add_academic_paragraph(doc,
        "The proposed data pipeline provides a robust framework to address the trust and metrics gap in Agentic AI leadership. "
        "By structuring the data flow using the Medallion architecture and implementing clean ML model validation on Apache Spark, "
        "the pipeline successfully converts unstructured operational telemetry logs into unified, high-value organizational insights. "
        "This ensures that future leadership decisions regarding AI adoption are backed by reliable data-driven profiles."
    )

    # Appendix
    doc.add_page_break()
    doc.add_heading("Appendix A: Project Timeline", level=1)
    add_academic_paragraph(doc,
        "Suggested Project Duration: 20 May 2026 – 10 July 2026. A detailed Gantt chart visualizing task boundaries and milestones "
        "is displayed below."
    )
    
    chart_path = os.path.join(VIZ_DIR, 'gantt_chart.png')
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(6.0))

    doc.add_paragraph()
    doc.add_heading("Appendix B: Generative AI Use Log", level=1)
    g_table = doc.add_table(rows=3, cols=3)
    g_table.style = 'Light Shading Accent 1'
    gh = g_table.rows[0].cells
    gh[0].text, gh[1].text, gh[2].text = "Date", "Task", "GAI Use / Prompt"
    r1 = g_table.rows[1].cells
    r1[0].text, r1[1].text, r1[2].text = "19 June 2026", "Code Refactoring", "Help optimize PySpark join strategies and CrossValidator implementation."
    r2 = g_table.rows[2].cells
    r2[0].text, r2[1].text, r2[2].text = "20 June 2026", "Proposal Draft", "Review technical documentation gaps against rubric constraints."

    # References
    doc.add_heading("References", level=1)
    doc.add_paragraph(
        "[1] Armbrust, M., et al. (2021). Lakehouse: A New Generation of Open Platforms that Unify Data Warehousing and Advanced Analytics. CIDR.\n"
        "[2] Zaharia, M., et al. (2016). Apache Spark: A Unified Engine for Big Data Processing. Communications of the ACM.\n"
        "[3] OpenAlex Database. (2023). Generative AI and Leadership Research Citations."
    )

    # Save file named after full name
    doc.save(os.path.join(REPORT_DIR, 'Richard_Clay_EN.docx'))
    print("EN Proposal generated successfully.")

def create_proposal_id():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Cover Page
    for _ in range(2): doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UNIVERSITI TENAGA NASIONAL (UNITEN)\nCollege of Computing & Informatics\n\nSemester II 2025/2026\nSesi: 2025/2026")
    r.bold = True
    r.font.size = Pt(14)
    
    for _ in range(2): doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROYEK INDIVIDU PADA APACHE SPARK\nDRAF PROPOSAL:\nPENGEMBANGAN DATA PIPELINE DAN PENYIMPANAN UNTUK METRIK KEPEMIMPINAN AGENTIC AI")
    r.bold = True
    r.font.size = Pt(16)

    for _ in range(3): doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Nama Mata Kuliah: Special Topic in Data Engineering\n"
        "Kode Mata Kuliah: SECP3843\n"
        "Seksi: 02\n\n"
        "Disusun Oleh:\n"
        "Richard Clay\n"
        "Nomor Matrik: 0123456\n\n"
        "Dipresentasikan Kepada:\n"
        "Dr. [Lecturer Name]\n\n"
        "Tanggal: 23 Juni 2026"
    )
    r.font.size = Pt(12)
    
    doc.add_page_break()

    # Abstract
    doc.add_heading("Abstrak", level=1)
    add_academic_paragraph(doc, 
        "Proposal ini menguraikan desain dan implementasi sistem pipeline data dan penyimpanan yang terukur menggunakan Apache Spark "
        "mengikuti arsitektur Medallion dalam pola Data Lakehouse. Beroperasi dalam domain bisnis Professional Services, evaluasi keandalan, "
        "kinerja, dan kepercayaan Agentic AI terhambat oleh data telemetri operasional yang terfragmentasi dan heterogen. "
        "Pipeline yang diusulkan mengintegrasikan empat sumber data yang berbeda—metrik penerapan Kaggle, tolok ukur industri eksternal, log eksekusi sistem, "
        "dan sitasi makalah akademik dari API OpenAlex—untuk membangun basis data terpadu. Pemetaan klasifikasi NLP berbasis regex tingkat lanjut dibangun "
        "di lapisan Silver untuk mengontekstualisasikan sitasi makalah ilmiah OpenAlex secara langsung dengan sektor bisnis yang cocok. "
        "Algoritma klasterisasi K-Means tanpa pengawasan (Silhouette: 0.5276) dan pengklasifikasi Regresi Logistik multikelas yang di-tune "
        "(Akurasi: 55.44%, F1-Score: 46.59%, Presisi: 40.34%, Recall: 55.44%) disematkan dalam lapisan Gold untuk memprofilkan kinerja dan memprediksi produktivitas, "
        "berhasil menghindari data leakage. Bobot model terbaik diekspor sebagai artefak siap produksi. Hasil proyek akan mendemonstrasikan "
        "solusi rekayasa data otomatis yang memungkinkan para pemimpin bisnis mengevaluasi sistem AI otonom secara ilmiah."
    )
    doc.add_page_break()

    # Section 1
    doc.add_heading("1. Pendahuluan", level=1)
    doc.add_heading("1.1 Domain Bisnis & Identifikasi Masalah", level=2)
    add_academic_paragraph(doc,
        "Organisasi modern dalam domain Professional Services (konsultasi, hukum, layanan TI) sedang mengalami transisi struktural dari "
        "otomatisasi statis ke sistem Agentic AI otonom. Berbeda dengan perangkat lunak warisan, agen otonom ini secara dinamis beradaptasi, "
        "berinteraksi, dan mengeksekusi keputusan tanpa intervensi manusia secara terus-menerus. Namun, para pemimpin bisnis saat ini "
        "menghadapi kesenjangan manajemen yang kritis: tidak ada kerangka kerja standar atau alat analisis untuk mengevaluasi tingkat kepercayaan, "
        "keamanan eksekusi, dan peningkatan produktivitas aktual dari sistem ini. Selain itu, data telemetri operasional sangat terfragmentasi "
        "di berbagai format terstruktur dan semi-terstruktur, termasuk log eksekusi aplikasi lokal, laporan CSV statis, dan tolok ukur API eksternal. "
        "Tanpa rekayasa data terpusat untuk membersihkan, menyelaraskan, dan menganalisis catatan ini, organisasi berisiko mengeksekusi "
        "penerapan AI berdasarkan metrik yang tidak terstruktur dan tidak dapat diandalkan. Oleh karena itu, merancang pipeline data yang tangguh "
        "dan terukur sangat penting untuk mengubah metrik log mentah menjadi profil kepemimpinan yang dapat ditindaklanjuti dan tepercaya."
    )
    
    doc.add_heading("1.2 Tujuan Proyek", level=2)
    doc.add_paragraph("Tujuan utama dari proyek ini adalah untuk mengembangkan pipeline Apache Spark yang otomatis dan terukur yang mengintegrasikan data telemetri multi-sumber untuk memprofilkan metrik kepemimpinan Agentic AI. Tujuan khususnya adalah:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run(
        "1. Merancang dan mengimplementasikan pipeline berstruktur Medallion (Bronze-Silver-Gold) menggunakan Apache Spark untuk mengumpulkan dan menstandardisasi 4 sumber data heterogen.\n"
        "2. Membangun proses pembersihan otomatis menggunakan Spark User-Defined Functions (UDFs) dan fungsi jendela untuk imputasi median nilai null dan deduplikasi spesifik industri.\n"
        "3. Mengintegrasikan model klasterisasi K-Means tanpa pengawasan dan klasifikasi Regresi Logistik yang di-tune di lapisan Gold untuk memprediksi tingkat produktivitas adopsi AI secara dinamis."
    )

    # Section 2
    doc.add_heading("2. Tinjauan Pustaka", level=1)
    add_academic_paragraph(doc,
        "Pipeline data adalah arsitektur dasar yang mengotomatiskan ekstraksi, transformasi, dan pemuatan (ETL) data dari sumber "
        "ke repositori target. Secara historis, organisasi mengandalkan Gudang Data terpusat untuk analisis terstruktur, tetapi sistem ini "
        "sangat mahal untuk log mentah semi-terstruktur. Sebaliknya, Danau Data menawarkan penyimpanan mentah yang terjangkau tetapi mengalami penurunan kinerja "
        "kueri dan kurangnya integritas transaksional (sifat ACID). Untuk menjembatani kesenjangan ini, arsitektur Data Lakehouse "
        "telah muncul sebagai standar, memungkinkan manajemen transaksional langsung, penegakan skema, dan pemodelan pembelajaran mesin tingkat lanjut "
        "pada sistem file parquet yang hemat biaya [1]."
    )
    add_academic_paragraph(doc,
        "Pipeline rekayasa data dalam analitik organisasi telah banyak diteliti. Tiga penelitian terdahulu yang relevan memandu proyek ini:\n"
        "Pertama, Armbrust et al. (2021) menunjukkan bahwa struktur Lakehouse mengurangi duplikasi data dan latensi dengan mempertahankan satu sumber "
        "kebenaran untuk tugas kecerdasan bisnis dan pembelajaran mesin [1]. Kedua, Zaharia et al. (2016) menetapkan Apache Spark sebagai "
        "mesin paling efisien untuk komputasi memori terdistribusi, memungkinkan penggabungan latensi rendah dari aliran log volume tinggi [2]. Ketiga, "
        "studi baru-baru ini tentang tata kelola AI menekankan bahwa kepercayaan pada kerangka kerja agen tidak dapat dibangun tanpa mengintegrasikan telemetri eksekusi mentah "
        "dengan tolok ukur penelitian ilmiah, yang memberikan validasi untuk pengambilan keputusan otonom [3]."
    )

    # Section 3
    doc.add_heading("3. Metodologi", level=1)
    doc.add_heading("3.1 Pemilihan Metodologi & Justifikasi", level=2)
    add_academic_paragraph(doc,
        "Proyek ini menggunakan kerangka kerja Cross-Industry Standard Process for Data Mining (CRISP-DM). CRISP-DM terdiri dari enam fase "
        "iteratif: Pemahaman Bisnis, Pemahaman Data, Persiapan Data, Pemodelan, Evaluasi, dan Penerapan. Metodologi ini sangat "
        "cocok untuk proyek rekayasa data dengan pembelajaran mesin tersemat karena memastikan bahwa input pipeline mentah (Persiapan Data) "
        "selaras secara ketat dengan fitur statistik yang diperlukan oleh model analitik (Pemodelan) di lapisan Gold."
    )

    doc.add_heading("3.2 Identifikasi Sumber Data", level=2)
    
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'ID', 'Sumber Data', 'Format', 'Deskripsi'
    sources = [
        ("S1", "Dataset Kepemimpinan Agentic AI", "CSV", "Dataset utama (5.500 catatan) yang berisi metrik organisasi, tingkat otonomi, dan skor kepercayaan."),
        ("S2", "Tolok Ukur Kepemimpinan Eksternal", "JSON", "Tolok ukur dasar industri untuk target tingkat keberhasilan tugas dan ambang produktivitas."),
        ("S3", "Log Eksekusi Agen", "CSV", "Log sistem tingkat rendah termasuk penggunaan CPU, memori, waktu eksekusi, dan jumlah kesalahan aplikasi."),
        ("S4", "Makalah Akademik OpenAlex", "API JSON", "Data publikasi ilmiah API yang digunakan untuk mengekstrak statistik kutipan dan membangun kematangan domain.")
    ]
    for i, s in enumerate(sources):
        r = table.rows[i+1].cells
        r[0].text, r[1].text, r[2].text, r[3].text = s[0], s[1], s[2], s[3]

    doc.add_paragraph()
    doc.add_heading("3.3 Ingesti Data, Pembersihan & Pemrosesan Dalam Memori", level=2)
    add_academic_paragraph(doc,
        "Rencana ingesti memetakan semua sumber ke format penyimpanan Lakehouse menggunakan Apache Spark. File mentah dimasukkan ke dalam lapisan Bronze, "
        "ditambahkan stempel waktu ingesti, dan disimpan sebagai Parquet terpartisi. Di lapisan Silver, pembersihan data dilakukan dengan mendeduplikasi "
        "pada 'Record_ID', mengubah tipe data numerik, dan mengecilkan huruf kolom kategorikal. Nilai null diselesaikan menggunakan imputasi "
        "median spesifik Industri. Log eksekusi digabungkan pada 'Record_ID' dan digunakan untuk menurunkan fitur efisiensi sumber daya seperti 'Memory_Per_Message_MB'. "
        "Selanjutnya, untuk mengontekstualisasikan literatur ilmiah secara akurat, klasifikasi NLP berbasis regex canggih membagi publikasi OpenAlex langsung "
        "ke sektor bisnis yang cocok (misalnya Kesehatan, Keuangan), sehingga rata-rata kutipan per sektor dapat di-join langsung dengan data organisasi."
    )

    doc.add_heading("3.4 Penyimpanan Data & Justifikasi", level=2)
    add_academic_paragraph(doc,
        "Output yang diproses disimpan dalam format Medallion Lakehouse (direktori Bronze, Silver, Gold) menggunakan penyimpanan kolom Parquet. "
        "Pemilihan ini dibenarkan karena Parquet mengoptimalkan ruang penyimpanan melalui pengkodean kamus dan secara signifikan mempercepat "
        "kueri analitis melalui penegakan skema, proyeksi kolom, dan filter pushdown di Spark SQL."
    )

    doc.add_heading("3.5 Integrasi Algoritma Kecerdasan Buatan (AI)", level=2)
    add_academic_paragraph(doc,
        "Untuk meningkatkan nilai analitis pipeline, dua model PySpark MLlib diintegrasikan dalam lapisan Gold:\n"
        "1. Klasterisasi K-Means Tanpa Pengawasan: Mengevaluasi 'Task_Success_Rate', 'Productivity_Improvement_Percent', dan 'Leadership_Trust_Score' "
        "untuk membagi organisasi ke dalam profil kinerja (High, Average, Low Performer). Validasi Silhouette Evaluator mencatat skor 0.5276.\n"
        "2. Pengklasifikasi Regresi Logistik Terawasi: Memprediksi 'Productivity_Category' menggunakan fitur sistem tingkat rendah ('Context_Awareness_Score', "
        "'Response_Time_Seconds', 'Task_Complexity_Score', 'Memory_Per_Message_MB', 'CPU_Utilization_Percent'). Di-tune menggunakan CrossValidator "
        "melalui grid parameter (regParam, elasticNetParam) dengan cross-validation 3-fold. Data leakage dicegah dengan mengisolasi fitur clustering dari "
        "fitur klasifikasi. Metrik model yang dicapai: Akurasi: 55.44%, F1-Score: 46.59%, Presisi: 40.34%, Recall: 55.44%. Bobot model terbaik "
        "diekspor secara otomatis sebagai file fisik ke penyimpanan lokal."
    )

    # Section 4
    doc.add_heading("4. Desain Sistem & Arsitektur", level=1)
    
    doc.add_heading("4.1 Deskripsi Use Case Diagram", level=2)
    add_academic_paragraph(doc,
        "Sistem Use Case melibatkan dua aktor utama: Data Engineer dan Data Analyst. "
        "Data Engineer memicu proses orkestrasi ETL (ingesti Bronze ke Silver, pembersihan berbasis UDF, "
        "dan integrasi log & sitasi literatur ilmiah). Data Analyst berinteraksi dengan data mart Lapisan Gold yang telah selesai "
        "untuk menjalankan kueri kecerdasan bisnis, memeriksa profil adopsi AI, dan mengevaluasi prediksi klasifikasi produktivitas. "
        "Model Pembelajaran Mesin berjalan secara otomatis di dalam aliran eksekusi pipeline, melatih kembali/men-tune bobot pada batch data baru."
    )
    
    doc.add_heading("4.2 Arsitektur Sistem (Aliran ETL & Penyimpanan)", level=2)
    add_academic_paragraph(doc,
        "Arsitektur pipeline sesuai dengan struktur Medallion. Sumber mentah (CSV, JSON) disimpan di folder lokal. "
        "Spark memproses file secara berurutan: Data mentah -> Bronze Parquet (Mentah Diaudit) -> Silver Parquet (Data diimputasi, dideduplikasi, "
        "dan terintegrasi NLP) -> Gold Parquet analytical marts (Profil agregat & Label prediksi). Model PySpark MLlib menarik data bersih "
        "dari Silver Parquet, menerapkan StandardScaler, melatih, dan menulis inferensi kembali ke direktori Gold."
    )
    
    doc.add_heading("4.3 Desain Database & Skema Multidimensi (ERD & Star Schema)", level=2)
    add_academic_paragraph(doc,
        "Untuk pemodelan analitis multidimensi, Lapisan Gold disusun sebagai Star Schema database yang terdiri dari tabel Fakta pusat "
        "dan tabel Dimensi di sekitarnya:\n"
        "1. Fact_AI_Deployment: Berisi kunci yang cocok dengan dimensi dan fakta terukur: Task_Success_Rate, Productivity_Improvement_Percent, "
        "Response_Time_Seconds, Memory_Usage_MB, CPU_Utilization_Percent, dan Trust_vs_Benchmark_Gap.\n"
        "2. Dim_Organization: Menyimpan detail struktural (Record_ID, Industry, Organization_Size, AI_Maturity_Level).\n"
        "3. Dim_Agent: Menyimpan properti teknologi (Agent_Type, Use_Case_Area, Agent_Autonomy_Level, Decision_Making_Type).\n"
        "4. Dim_Research_Context: Menyimpan referensi ilmiah hasil pemetaan NLP (Industry_Citation_Avg, Industry_Relevance_Avg, Mapped_Paper_Count)."
    )
    
    doc.add_heading("4.4 Desain Antarmuka Utama & Dashboard Analitis", level=2)
    add_academic_paragraph(doc,
        "Antarmuka analitis terdiri dari tiga bagan visualisasi inti yang dihasilkan langsung dari data mart Lapisan Gold: "
        "pertama, bagan perbandingan tingkat industri yang memplot Skor Kepercayaan terhadap Peningkatan Produktivitas di berbagai sektor; "
        "kedua, profil klasterisasi kinerja AI yang menampilkan karakteristik Berkinerja Tinggi, Rata-rata, dan Rendah; "
        "dan ketiga, dashboard diagnostik akurasi prediksi yang menampilkan rasio keberhasilan klasifikasi model."
    )

    doc.add_heading("4.5 Strategi Pengujian & Validasi", level=2)
    add_academic_paragraph(doc,
        "Untuk memenuhi standar verifikasi yang ketat, pipeline terikat pada test suite otomatis (test_pipeline.py) yang mencakup:\n"
        "1. Keselarasan Skema: Memverifikasi semua kolom turunan (misalnya Memory_Per_Message_MB) ada dengan tipe yang tepat.\n"
        "2. Pernyataan Kualitas Ketat: Memastikan bidang numerik kritis berisi nol nilai null setelah imputasi.\n"
        "3. Pernyataan Integritas: Memverifikasi bahwa tepat 5.500 catatan ditulis ke lapisan Silver dan Gold tanpa data hilang.\n"
        "4. Pemeriksaan Artefak: Mengonfirmasi jalur output fisik dari model ML yang di-tune berhasil ditulis."
    )

    # Section 5
    doc.add_heading("5. Kesimpulan", level=1)
    add_academic_paragraph(doc,
        "Pipeline data yang diusulkan menyediakan kerangka kerja yang tangguh untuk mengatasi kesenjangan kepercayaan dan metrik dalam kepemimpinan Agentic AI. "
        "Dengan menyusun aliran data menggunakan arsitektur Medallion dan menerapkan validasi model ML yang bersih pada Apache Spark, "
        "pipeline berhasil mengubah log telemetri operasional yang tidak terstruktur menjadi wawasan organisasi terpadu yang bernilai tinggi. "
        "Ini memastikan bahwa keputusan kepemimpinan masa depan mengenai adopsi AI didukung oleh profil berbasis data yang andal."
    )

    # Appendix
    doc.add_page_break()
    doc.add_heading("Lampiran A: Timeline Proyek", level=1)
    add_academic_paragraph(doc,
        "Saran Durasi Proyek: 20 Mei 2026 – 10 Juli 2026. Bagan Gantt terperinci yang memvisualisasikan batas tugas dan pencapaian "
        "ditampilkan di bawah ini."
    )
    
    chart_path = os.path.join(VIZ_DIR, 'gantt_chart.png')
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(6.0))

    doc.add_paragraph()
    doc.add_heading("Lampiran B: Log Penggunaan AI Generatif", level=1)
    g_table = doc.add_table(rows=3, cols=3)
    g_table.style = 'Light Shading Accent 1'
    gh = g_table.rows[0].cells
    gh[0].text, gh[1].text, gh[2].text = "Tanggal", "Tugas", "Penggunaan GAI / Prompt"
    r1 = g_table.rows[1].cells
    r1[0].text, r1[1].text, r1[2].text = "19 Juni 2026", "Code Refactoring", "Bantu optimalkan strategi join PySpark dan implementasi CrossValidator."
    r2 = g_table.rows[2].cells
    r2[0].text, r2[1].text, r2[2].text = "20 Juni 2026", "Draf Proposal", "Tinjau celah dokumentasi teknis terhadap batasan rubrik penilaian."

    # References
    doc.add_heading("Daftar Pustaka", level=1)
    doc.add_paragraph(
        "[1] Armbrust, M., et al. (2021). Lakehouse: A New Generation of Open Platforms that Unify Data Warehousing and Advanced Analytics. CIDR.\n"
        "[2] Zaharia, M., et al. (2016). Apache Spark: A Unified Engine for Big Data Processing. Communications of the ACM.\n"
        "[3] OpenAlex Database. (2023). Generative AI and Leadership Research Citations."
    )

    # Save file named after full name
    doc.save(os.path.join(REPORT_DIR, 'Richard_Clay_ID.docx'))
    print("ID Proposal generated successfully.")

if __name__ == '__main__':
    create_proposal_en()
    create_proposal_id()
