# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REPORT_DIR = os.path.abspath(os.path.join(BASE_DIR, '..', '07_Reports'))
VIZ_DIR = os.path.abspath(os.path.join(BASE_DIR, '..', '06_Visualization'))
TEMPLATE_PATH = r"C:\Users\richa\Documents\Sem.6\DatEng_Individual_Project\Materials\PROPOSAL_individual project.docx"
CHART_UC = os.path.join(VIZ_DIR, 'use_case_diagram.png')
CHART_ERD = os.path.join(VIZ_DIR, 'star_schema_erd.png')

BLACK = RGBColor(0, 0, 0)
TNR = "Times New Roman"

def set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, line_spacing=1.5, first_line_indent=0.5):
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent > 0:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)

def apply_font(p, size=12, bold=False, italic=False, color=BLACK):
    for run in p.runs:
        run.font.name = TNR
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = color

def add_heading_1(doc, text):
    p = doc.add_paragraph(style='Heading 1')
    p.clear()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=12, line_spacing=1.15, first_line_indent=0)
    run = p.add_run(text)
    run.font.name = TNR
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = BLACK
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    p.clear()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, line_spacing=1.15, first_line_indent=0)
    run = p.add_run(text)
    run.font.name = TNR
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = BLACK
    return p

def add_normal_para(doc, text, first_line_indent=0.5):
    p = doc.add_paragraph()
    set_para_format(p, first_line_indent=first_line_indent)
    p.add_run(text)
    apply_font(p)
    return p

def make_toc(doc, title, is_en=True):
    """Build a static Table of Contents listing all sections and subsections."""
    doc.add_page_break()
    # Title
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(title)
    run.font.name = TNR
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = BLACK

    if is_en:
        entries = [
            ("Abstract", 1), ("1. Introduction", 1), ("   1.1 Business Domain & Problem Identification", 2),
            ("   1.2 Objectives", 2), ("   1.3 Risk Analysis & Mitigation Plan", 2),
            ("2. Literature Review", 1), ("3. Methodology", 1),
            ("   3.1 Methodology Selection & Justification", 2),
            ("   3.2 Data Source Identification", 2),
            ("   3.3 Data Ingestion, Cleansing & Processing", 2),
            ("   3.4 Data Storage & Justification", 2),
            ("   3.5 AI Algorithm Integration", 2),
            ("4. System Design & Architecture", 1),
            ("   4.1 Use Case Diagram Description", 2),
            ("   4.2 System Architecture (ETL & Storage Flow)", 2),
            ("   4.3 Database & Multidimensional Schema (Star Schema)", 2),
            ("   4.4 Main Interface & Analytical Dashboard Design", 2),
            ("   4.5 Testing & Validation Strategy", 2),
            ("5. Conclusion", 1), ("Appendix A: Project Timeline", 1),
            ("Appendix B: Generative AI Use Log", 1), ("References", 1),
        ]
    else:
        entries = [
            ("Abstrak", 1), ("1. Pendahuluan", 1), ("   1.1 Domain Bisnis & Identifikasi Masalah", 2),
            ("   1.2 Tujuan Proyek", 2), ("   1.3 Analisis Risiko & Rencana Mitigasi", 2),
            ("2. Tinjauan Pustaka", 1), ("3. Metodologi", 1),
            ("   3.1 Pemilihan Metodologi & Justifikasi", 2),
            ("   3.2 Identifikasi Sumber Data", 2),
            ("   3.3 Ingesti Data, Pembersihan & Pemrosesan", 2),
            ("   3.4 Penyimpanan Data & Justifikasi", 2),
            ("   3.5 Integrasi Algoritma Kecerdasan Buatan", 2),
            ("4. Desain Sistem & Arsitektur", 1),
            ("   4.1 Deskripsi Use Case Diagram", 2),
            ("   4.2 Arsitektur Sistem (Aliran ETL & Penyimpanan)", 2),
            ("   4.3 Desain Database & Skema Multidimensi", 2),
            ("   4.4 Antarmuka Utama & Dashboard Analitis", 2),
            ("   4.5 Strategi Pengujian & Validasi", 2),
            ("5. Kesimpulan", 1), ("Lampiran A: Timeline Proyek", 1),
            ("Lampiran B: Log Penggunaan AI Generatif", 1), ("Daftar Pustaka", 1),
        ]
            ("   1.1 Business Domain & Problem Identification", 2),
            ("   1.2 Objectives", 2),
            ("   1.3 Risk Analysis & Mitigation Plan", 2),
            ("2. Literature Review", 1),
            ("3. Methodology", 1),
            ("   3.1 Methodology Selection & Justification", 2),
            ("   3.2 Data Source Identification", 2),
            ("   3.3 Data Ingestion, Cleansing & Processing", 2),
            ("   3.4 Data Storage & Justification", 2),
            ("   3.5 AI Algorithm Integration", 2),
            ("4. System Design & Architecture", 1),
            ("   4.1 Use Case Diagram Description", 2),
            ("   4.2 System Architecture (ETL & Storage Flow)", 2),
            ("   4.3 Database & Multidimensional Schema (Star Schema)", 2),
            ("   4.4 Main Interface & Analytical Dashboard Design", 2),
            ("   4.5 Testing & Validation Strategy", 2),
            ("5. Conclusion", 1),
            ("Appendix A: Project Timeline", 1),
            ("Appendix B: Generative AI Use Log", 1),
            ("References", 1),
        ]
    else:
        entries = [
            ("Abstrak", 1),
            ("1. Pendahuluan", 1),
            ("   1.1 Domain Bisnis & Identifikasi Masalah", 2),
            ("   1.2 Tujuan Proyek", 2),
            ("   1.3 Analisis Risiko & Rencana Mitigasi", 2),
            ("2. Tinjauan Pustaka", 1),
            ("3. Metodologi", 1),
            ("   3.1 Pemilihan Metodologi & Justifikasi", 2),
            ("   3.2 Identifikasi Sumber Data", 2),
            ("   3.3 Ingesti Data, Pembersihan & Pemrosesan", 2),
            ("   3.4 Penyimpanan Data & Justifikasi", 2),
            ("   3.5 Integrasi Algoritma Kecerdasan Buatan", 2),
            ("4. Desain Sistem & Arsitektur", 1),
            ("   4.1 Deskripsi Use Case Diagram", 2),
            ("   4.2 Arsitektur Sistem (Aliran ETL & Penyimpanan)", 2),
            ("   4.3 Desain Database & Skema Multidimensi", 2),
            ("   4.4 Antarmuka Utama & Dashboard Analitis", 2),
            ("   4.5 Strategi Pengujian & Validasi", 2),
            ("5. Kesimpulan", 1),
            ("Lampiran A: Timeline Proyek", 1),
            ("Lampiran B: Log Penggunaan AI Generatif", 1),
            ("Daftar Pustaka", 1),
        ]

    for entry_text, level in entries:
        p_entry = doc.add_paragraph()
        p_entry.alignment = 0  # LEFT
        p_entry.paragraph_format.space_after = Pt(2)
        p_entry.paragraph_format.space_before = Pt(2)
        p_entry.paragraph_format.line_spacing = 1.2
        
        indent = Inches(0.3) if level == 2 else Inches(0)
        p_entry.paragraph_format.left_indent = indent

        run = p_entry.add_run(entry_text)
        run.font.name = TNR
        run.font.size = Pt(12) if level == 1 else Pt(11)
        run.bold = (level == 1)
        run.font.color.rgb = BLACK

def append_en_content(doc):
    make_toc(doc, "Table of Contents", is_en=True)

    # Abstract
    doc.add_page_break()
    add_heading_1(doc, "Abstract")
    add_normal_para(doc,
        "This project builds a data pipeline to measure how well Agentic AI systems perform in enterprise environments. "
        "Currently, decision-makers lack hard numbers to trust these AI agents, relying instead on scattered logs "
        "and subjective estimates. To solve this, we used Apache Spark to construct a Medallion architecture (Bronze, Silver, Gold) "
        "pipeline. We integrated four datasets: a CSV with 5,500 company records, JSON industry benchmarks, system "
        "execution logs, and OpenAlex academic citations. The pipeline cleans the data, handles missing values using "
        "industry medians, and maps academic papers to business sectors using NLP. In the Gold layer, we use K-Means "
        "clustering (Silhouette score 0.5276) to group companies by performance, and a tuned Logistic Regression model "
        "(56.01% accuracy) to predict productivity. We isolated the clustering features from the prediction model to prevent "
        "data leakage. The final model is serialized and ready for deployment, providing leadership with actionable, "
        "validated numbers instead of guesswork.")

    # Chapter 1
    doc.add_page_break()
    add_heading_1(doc, "1. Introduction")

    # Introductory paragraphs (like CONTOH REPORT style)
    add_normal_para(doc,
        "Agentic AI is changing how businesses operate. Unlike rule-based bots that just follow scripts, "
        "these systems plan, use tools, and adapt on their own. Companies deploy them for customer service, "
        "supply chains, and IT operations to cut costs and speed up decisions. But there is a catch: proving "
        "that these agents actually work is difficult. Managers need hard numbers, but the data is usually scattered.")

    add_normal_para(doc,
        "System logs are stuck in text files, company metrics are in CSVs, and industry benchmarks are hidden "
        "behind APIs. Without a unified pipeline to clean and join these sources, companies end up guessing "
        "if their AI is actually performing well. Decision-makers cannot compare results, spot anomalies, or "
        "trust the AI without a proper data infrastructure.")

    add_normal_para(doc,
        "This project fixes that measurement problem. We built a Spark-based Medallion pipeline that "
        "turns messy logs into clear performance profiles. The Bronze layer ingests the raw data. The Silver "
        "layer cleans it, handles missing values, and uses NLP to map academic citations to business sectors. "
        "Finally, the Gold layer applies machine learning to predict performance. The end result is a reliable "
        "dataset that helps leaders make data-backed decisions about their AI systems.")

    add_heading_2(doc, "1.1 Business Domain & Problem Identification")
    add_normal_para(doc,
        "The business domain selected for this project is Professional Services, a sector that encompasses "
        "consulting firms, legal practices, IT service providers, and financial advisory organisations. "
        "These enterprises are characterized by knowledge-intensive operations where human expertise is "
        "augmented by increasingly autonomous technological systems. Within this domain, Agentic AI "
        "applications are being deployed for tasks such as automated document review, contract analysis, "
        "client communication triage, data-driven recommendation generation, and intelligent process "
        "automation. The performance of these systems directly impacts client satisfaction, operational "
        "efficiency, and regulatory compliance, making reliable evaluation frameworks a critical "
        "organizational requirement.")

    add_normal_para(doc,
        "Despite the clear need for performance evaluation, the current state of practice reveals a "
        "fundamental problem: there is no standardised data engineering framework that unifies the "
        "fragmented telemetry generated by Agentic AI systems. Operational data is produced in "
        "multiple disconnected formats\u2014structured CSV exports from deployment dashboards, "
        "semi-structured JSON responses from external benchmarking APIs, raw application log files "
        "recording system-level execution metrics, and academic citation data from public research "
        "databases. Each of these sources uses a different schema, naming convention, and data "
        "quality profile. Without systematic integration, leadership teams cannot perform cross-source "
        "analysis, aggregate performance trends, or correlate operational efficiency with external "
        "standards. This severely limits the ability of organisations to make informed, evidence-based "
        "decisions about their AI strategy and investments.")

    add_heading_2(doc, "1.2 Objectives")
    add_normal_para(doc,
        "The primary objective of this project is to design and implement an automated, scalable data "
        "engineering pipeline using Apache Spark that integrates multi-source telemetry data to produce "
        "reliable Agentic AI leadership performance profiles. The specific objectives are as follows:",
        first_line_indent=0)

    objs = [
        "1. To build a Medallion-structured (Bronze, Silver, Gold) data pipeline using Apache Spark that "
           "ingests, standardises, and integrates four heterogeneous data sources\u2014a primary CSV dataset "
           "containing 5,500 organisational records, an external industry benchmark JSON file, system "
           "execution logs in CSV format, and academic citation data from the OpenAlex API.",

        "2. To construct an automated data cleansing pipeline that employs Spark User-Defined Functions "
           "and window-based aggregation operations to perform industry-specific median imputation for "
           "null values, deduplication, type standardisation, derived feature engineering, and NLP-based "
           "contextual mapping of scientific paper citations to matching business sectors.",

        "3. To deploy and validate both an unsupervised K-Means clustering model and a tuned multiclass "
           "Logistic Regression classifier in the Gold processing layer, ensuring data isolation between "
           "clustering features and classification inputs to prevent data leakage, and exporting the "
           "best-performing model as a production-ready serialised artifact."
    ]
    for obj in objs:
        add_normal_para(doc, obj)

    add_heading_2(doc, "1.3 Project Risk Analysis & Mitigation Plan")
    add_normal_para(doc,
        "The implementation of a data engineering pipeline for Agentic AI evaluation carries several "
        "inherent risks. Identifying these risks early ensures that appropriate mitigation strategies "
        "can be integrated into the system architecture and project schedule.")
        
    r_table = doc.add_table(rows=4, cols=3)
    r_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_hdr = r_table.rows[0].cells
    r_hdr[0].text, r_hdr[1].text, r_hdr[2].text = "Identified Risk", "Impact / Probability", "Mitigation Strategy"
    for c in r_hdr:
        apply_font(c.paragraphs[0], bold=True)

    risks = [
        ("Data Quality & Schema Drift", "Impact: High\nProbability: Medium", 
         "Implement median-based imputation for missing numericals and validate schema integrity using automated PySpark unit tests before loading to the Gold layer."),
        ("Tool & Compute Limitations", "Impact: High\nProbability: High", 
         "The environment has limited RAM (1.9GB). Mitigated by tuning Spark partition counts, avoiding Python UDFs to prevent memory serialization overhead, and writing in batch chunks."),
        ("Stakeholder Alignment", "Impact: Medium\nProbability: Low", 
         "Ensuring ML outputs align with business KPIs by validating unsupervised clustering results (Silhouette Score) against leadership trust metrics.")
    ]
    for i, (r_name, r_imp, r_mit) in enumerate(risks):
        row = r_table.rows[i+1].cells
        row[0].text, row[1].text, row[2].text = r_name, r_imp, r_mit
        for c in row:
            apply_font(c.paragraphs[0])
            
    doc.add_paragraph()

    # Chapter 2
    doc.add_page_break()
    add_heading_1(doc, "2. Literature Review")
    add_normal_para(doc,
        "A solid data pipeline is essential for handling large volumes of enterprise data. Traditional "
        "Data Warehouses are too rigid for semi-structured logs, while Data Lakes lack transactional safety. "
        "The Data Lakehouse architecture solves this by combining the flexibility of a lake with the reliability "
        "of a warehouse. It uses columnar formats like Parquet and engines like Apache Spark to process data "
        "and run machine learning workloads directly (Armbrust et al., 2021).")

    add_normal_para(doc,
        "Our approach builds on three key findings. First, Lakehouse architecture reduces data duplication "
        "and simplifies pipelines (Armbrust et al., 2021). Second, Spark's in-memory processing is ideal "
        "for joining massive, complex datasets (Zaharia et al., 2016)\u2014exactly what we need for Agentic "
        "AI logs. Finally, recent AI governance research stresses that combining system logs with external "
        "baselines, like OpenAlex citations, is crucial for evaluating the real-world trustworthiness of "
        "autonomous agents (Priem et al., 2022). These papers define our architectural choices.")

    # Chapter 3
    doc.add_page_break()
    add_heading_1(doc, "3. Methodology")

    # Introductory paragraphs (like CONTOH REPORT style)
    add_normal_para(doc,
        "This section outlines how we built and tested the data pipeline. A clear methodology ensures "
        "the project hits its goals and can be reproduced later. We cover the framework choice, the four "
        "data sources, the ETL process, storage decisions, and the machine learning implementation.")

    add_normal_para(doc,
        "We use the Medallion Architecture pattern. The Bronze layer stores raw data unchanged. The "
        "Silver layer cleans and integrates it. The Gold layer applies analytics and machine learning. "
        "This setup keeps the code clean, handles bad data incrementally, and is standard practice in "
        "the industry. We managed the overall project using the CRISP-DM framework.")

    add_heading_2(doc, "3.1 Methodology Selection & Justification")
    add_normal_para(doc,
        "This project adopts the Cross-Industry Standard Process for Data Mining (CRISP-DM) as its "
        "primary methodological framework. CRISP-DM was selected because it is the most widely used "
        "and mature framework for data mining and analytics projects, providing a structured yet "
        "flexible lifecycle that can accommodate both data engineering and machine learning components. "
        "The framework consists of six iterative phases: Business Understanding, Data Understanding, "
        "Data Preparation, Modeling, Evaluation, and Deployment. For this project, the Business "
        "Understanding phase corresponds to the problem formulation described in Section 1, the Data "
        "Understanding and Data Preparation phases map to the Bronze and Silver layer processing steps, "
        "and the Modeling and Evaluation phases correspond to the Gold layer analytics. The iterative "
        "nature of CRISP-DM is particularly valuable because it allows the data preparation and "
        "modeling stages to inform each other\u2014an essential capability when integrating diverse data "
        "sources with different schemas and quality profiles.")

    add_heading_2(doc, "3.2 Data Source Identification")
    add_normal_para(doc,
        "This project utilises four distinct data sources to capture the full operational context of "
        "Agentic AI deployments. The sources were selected to represent the three main categories of "
        "data that organisations typically encounter: internal structured operational records (S1), "
        "external benchmarking standards (S2), granular system execution telemetry (S3), and academic "
        "research context (S4). The following table summarises each source, its format, and its "
        "role in the pipeline. Detailed column-level descriptions are provided in the subsequent "
        "paragraphs to support readers in understanding the schema that flows through each Medallion layer.")

    # Data source table
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'ID', 'Data Source', 'Format', 'Description'
    for c in hdr:
        apply_font(c.paragraphs[0], bold=True)

    sources = [
        ("S1", "Agentic AI Leadership Dataset", "CSV",
         "Primary operational dataset containing 5,500 records of organisational metrics, autonomy levels, task success rates, and leadership trust scores collected from simulated Agentic AI deployments across multiple industries."),
        ("S2", "External Leadership Benchmarks", "JSON",
         "Industry baseline benchmarks providing reference target values for task success rates and productivity thresholds, used to compute comparative performance gaps."),
        ("S3", "Agent Execution Logs", "CSV",
         "Granular system-level execution logs recording CPU utilisation percentages, memory consumption in MB, execution timestamps, response times in seconds, and application error counts per organisational record."),
        ("S4", "OpenAlex Academic Publications", "API JSON",
         "Academic metadata from the OpenAlex open-access database containing 25 scientific publication records with citation counts, publication years, and title keywords used to calculate domain-specific research maturity scores.")
    ]
    for i, s in enumerate(sources):
        r = table.rows[i+1].cells
        r[0].text, r[1].text, r[2].text, r[3].text = s[0], s[1], s[2], s[3]
        for c in r:
            apply_font(c.paragraphs[0])

    doc.add_paragraph()

    # Detailed CSV dataset description
    add_normal_para(doc,
        "Source S1, the Agentic AI Leadership Dataset, is the primary input to the pipeline. "
        "This dataset is stored in CSV (Comma-Separated Values) format and contains 5,500 records, "
        "each representing a distinct organisational unit that has deployed an Agentic AI system. "
        "The dataset includes 15 columns that capture organisational context, agent characteristics, "
        "and operational performance metrics. The organisational context columns include Record_ID "
        "(a unique alphanumeric identifier for each record), Organization_Name (the name of the "
        "enterprise deploying the AI system), Industry (the business sector classification such as "
        "Finance, Healthcare, or Education), Organization_Size (categorised as Small, Medium, or "
        "Large), and AI_Maturity_Level (a categorical variable indicating the organisation\u2019s "
        "experience level with AI, rated as Beginner, Intermediate, or Advanced).")

    add_normal_para(doc,
        "The agent configuration columns describe the characteristics of the deployed AI system "
        "and include Agent_Type (the framework type such as LangChain, AutoGPT, or CrewAI), "
        "Use_Case_Area (the functional domain such as Customer Support, Data Analysis, or Code "
        "Generation), Agent_Autonomy_Level (a categorical rating of the agent\u2019s operational "
        "independence), Decision_Making_Type (the algorithmic approach to decision-making such "
        "as Rule-based, LLM-based, or Hybrid), and Context_Awareness_Score (a numerical score "
        "between 0 and 100 measuring the agent\u2019s ability to perceive and interpret its operational "
        "context). The performance metric columns include Task_Success_Rate (the percentage of "
        "assigned tasks completed successfully, ranging from 50% to 100%), "
        "Productivity_Improvement_Percent (the measured improvement in operational output since "
        "deployment), Response_Time_Seconds (the average time taken by the agent to respond to "
        "queries or triggers), Complexity_Score (a normalised measure of task difficulty), and "
        "Leadership_Trust_Score (a weighted composite index reflecting executive confidence in "
        "the system\u2019s reliability). These 15 columns collectively provide a rich, multi-dimensional "
        "view of each organisation\u2019s experience with Agentic AI, enabling both clustering-based "
        "profiling and classification-based prediction within the Gold layer analytics.")

    add_heading_2(doc, "3.3 Data Ingestion, Cleansing & In-Memory Processing")
    add_normal_para(doc,
        "The data ingestion process begins with the Bronze layer, where the four source files are "
        "loaded into Apache Spark DataFrames using format-specific readers. CSV files are read with "
        "header inference and schema enforcement enabled, JSON files are parsed with multi-line "
        "support for nested structures, and the OpenAlex API response is consumed as a standard "
        "JSON file. Each ingested dataset is enriched with a metadata field recording the ingestion "
        "timestamp and source identifier before being written to partitioned Parquet files in the "
        "Bronze output directory. This raw storage layer serves as a permanent, unmodified archive "
        "that allows data lineage tracing and reprocessing if required.")

    add_normal_para(doc,
        "In the Silver layer, data cleansing operations are applied systematically. First, records "
        "are deduplicated by identifying and removing duplicate entries based on the Record_ID "
        "field. Second, all numeric columns are cast to their appropriate Spark data types to ensure "
        "computational consistency. Third, categorical text fields are standardised to lowercase "
        "to eliminate case-sensitivity issues during join operations. Fourth, missing numerical "
        "values are imputed using the median value calculated for each industry group; this approach "
        "preserves industry-specific distributional characteristics rather than applying a global "
        "imputation that could bias results. Execution log data is joined to the main dataset on "
        "Record_ID, enabling the derivation of efficiency features such as Memory_Per_Message_MB "
        "(calculated as total memory consumed divided by total requests processed). Finally, "
        "OpenAlex paper titles are processed through a regex-based classification script that "
        "maps each publication to a matching business sector, and the average citation count per "
        "sector is joined to the organisational records to contextualise scientific maturity.")

    add_heading_2(doc, "3.4 Data Storage & Justification")
    add_normal_para(doc,
        "All processed data is stored in Apache Parquet columnar format across the three Medallion "
        "directory layers: Bronze (raw archives), Silver (cleansed data), and Gold (analytical "
        "outputs). Parquet was selected as the storage format for several technical reasons. Its "
        "columnar storage layout enables predicate pushdown and column projection during Spark SQL "
        "queries, which significantly reduces I/O by reading only the columns needed for a given "
        "operation. Dictionary encoding further compresses repetitive string values, reducing storage "
        "footprint. Schema enforcement ensures that data written to each layer conforms to the "
        "expected structure, preventing silent corruption from schema drift. These characteristics "
        "make Parquet an industry-standard format for Lakehouse implementations and align with "
        "the architectural recommendations established in the literature review.")

    add_heading_2(doc, "3.5 Artificial Intelligence (AI) Algorithm Integration")
    add_normal_para(doc,
        "To deliver analytical value beyond basic aggregation, the Gold layer integrates two "
        "machine learning models from the PySpark MLlib library. These models operate on the "
        "cleaned and unified dataset produced by the Silver layer and are designed to generate "
        "actionable insights for business decision-makers.")

    add_normal_para(doc,
        "The first model is an unsupervised K-Means clustering algorithm configured to partition "
        "the 5,500 organisational records into three distinct performance tiers: High Performer, "
        "Average Performer, and Low Performer. The clustering is performed on three normalised "
        "features: Task_Success_Rate, Productivity_Improvement_Percent, and Leadership_Trust_Score. "
        "These features are first assembled into a feature vector using VectorAssembler, then "
        "standardised using StandardScaler to ensure equal weighting during distance computation. "
        "The K-Means algorithm is configured with k=3 clusters and a maximum of 20 iterations. "
        "Model quality is evaluated using the Silhouette score, which measures how similar each "
        "point is to its own cluster compared to other clusters. The achieved Silhouette score "
        "is 0.5276, indicating moderate cluster separation and confirming that the three identified "
        "profiles are distinguishable.")

    add_normal_para(doc,
        "The second model is a supervised multiclass Logistic Regression classifier that predicts "
        "the Productivity_Category (High, Average, or Low) based exclusively on system-level "
        "operational features: Context_Awareness_Score, Response_Time_Seconds, Task_Complexity_Score, "
        "Memory_Per_Message_MB, and CPU_Utilization_Percent. These features are intentionally "
        "selected to exclude the clustering metrics, thereby preventing data leakage. The model is "
        "tuned using a CrossValidator with 3-fold cross-validation over a parameter grid that "
        "searches combinations of regParam (regularisation strength) and elasticNetParam (mixing "
        "ratio between L1 and L2 regularisation). The tuned model achieves an accuracy of 56.01%, "
        "a weighted F1-score of 50.29%, a precision of 52.08%, and a recall of 56.01%. The best "
        "model pipeline, including the trained weights and scaler parameters, is serialised and "
        "exported to a physical file in the Gold output directory for production deployment.")

    # Chapter 4
    doc.add_page_break()
    add_heading_1(doc, "4. System Design & Architecture")

    # Long Introduction Paragraphs
    add_normal_para(doc,
        "The system architecture for this project is engineered to handle the end-to-end lifecycle of Agentic "
        "AI telemetry data. Moving from raw, fragmented data sources into a structured multidimensional schema "
        "requires a robust, fault-tolerant design. Apache Spark serves as the core computational engine, chosen "
        "specifically for its distributed in-memory processing capabilities which eliminate the Out-Of-Memory "
        "bottlenecks typically encountered with single-node tools like Pandas when joining large execution logs "
        "with external JSON API responses.")

    add_normal_para(doc,
        "This chapter details the structural design of the pipeline, divided into five core components. "
        "It begins with the Use Case diagram outlining user interactions, followed by the ETL data flow "
        "within the Medallion Lakehouse paradigm. Subsequent sections detail the Star Schema logical database "
        "design, the analytical dashboard interface intended for leadership consumption, and the rigorous "
        "unit testing strategy employed to guarantee zero data leakage between the K-Means clustering and "
        "Logistic Regression phases.")

    add_heading_2(doc, "4.1 Use Case Diagram Description")
    add_normal_para(doc,
        "The System Use Case involves two main actors: the Data Engineer and the Data Analyst. "
        "The Data Engineer is responsible for initiating the pipeline execution, monitoring "
        "processing status across all three Medallion layers, inspecting log files for errors, "
        "and triggering reprocessing if data quality issues are detected. The Data Analyst "
        "interacts with the Gold layer outputs to perform business intelligence queries, examine "
        "clustering profiles, evaluate classification predictions, and generate visualisation "
        "dashboards for leadership reporting. The machine learning models execute automatically "
        "as part of the pipeline orchestration, retraining on each fresh data batch and updating "
        "the serialised model artifact.")
        
    if os.path.exists(CHART_UC):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(CHART_UC, width=Inches(5.0))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.add_run("Figure 4.1.1: Use Case Diagram of the Lakehouse System").italic = True

    add_heading_2(doc, "4.2 System Architecture (ETL & Storage Flow)")
    add_normal_para(doc,
        "The pipeline architecture follows a modular Medallion design. Raw datasets are stored "
        "in a local folder structure and are processed sequentially by Apache Spark. The Bronze "
        "layer produces audited Parquet archives, the Silver layer generates cleansed and "
        "feature-engineered Parquet outputs, and the Gold layer produces aggregated analytical "
        "marts alongside machine learning predictions and the serialised model artifact. Each layer "
        "operates as an independent Spark job with its own input validation, transformation logic, "
        "and output verification steps. Logging is performed at each stage to provide operational "
        "visibility into pipeline execution.")

    add_heading_2(doc, "4.3 Database & Multidimensional Schema (Star Schema)")
    add_normal_para(doc,
        "For multidimensional analytical modelling, the Gold Layer is structured as a Star Schema "
        "database. The central fact table, Fact_AI_Deployment, contains foreign key references to "
        "three dimension tables: Dim_Organization (organisational attributes), Dim_Agent (AI agent "
        "configuration details), and Dim_Research_Context (NLP-mapped scientific citation metrics). "
        "This dimensional design enables business analysts to query performance metrics across "
        "multiple analytical axes\u2014for example, comparing average trust scores by industry, "
        "agent type, or research citation density\u2014without requiring complex multi-table joins "
        "at query runtime.")
        
    if os.path.exists(CHART_ERD):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(CHART_ERD, width=Inches(5.5))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.add_run("Figure 4.3.1: Star Schema ERD Diagram").italic = True

    add_heading_2(doc, "4.4 Main Interface & Analytical Dashboard Design")
    add_normal_para(doc,
        "The analytical interface is designed as three visualisation dashboards. The first dashboard "
        "displays an industry-level comparison chart plotting Leadership Trust Score against "
        "Productivity Improvement across all sectors. The second dashboard shows clustering profile "
        "distributions, enabling users to identify the characteristics of High, Average, and Low "
        "Performer groups. The third dashboard presents the classification model\u2019s prediction "
        "accuracy metrics, confusion matrix statistics, and per-class performance breakdowns. "
        "These dashboards are intended to be implemented in a business intelligence tool such as "
        "Power BI or Tableau, connecting directly to the Gold layer Parquet files.")

    add_heading_2(doc, "4.5 Testing & Validation Strategy")
    add_normal_para(doc,
        "To ensure data quality and pipeline reliability, an automated test suite is executed after "
        "each pipeline run. The suite verifies four aspects: schema conformance (checking that all "
        "expected derived columns exist with correct data types), null-value assertions (verifying "
        "that critical fields contain zero null values after Silver layer imputation), record "
        "integrity (confirming that exactly 5,500 records are preserved from Bronze through Gold), "
        "and model artifact validation (checking that the serialised model file exists at the "
        "expected output path). All tests run on PySpark DataFrames to avoid memory limitations "
        "associated with pandas operations on large datasets.")

    # Chapter 5
    doc.add_page_break()
    add_heading_1(doc, "5. Conclusion")
    add_normal_para(doc,
        "This proposal has presented a comprehensive design for a scalable, automated data "
        "engineering pipeline that addresses the critical gap in Agentic AI performance evaluation. "
        "By implementing the Medallion architecture on Apache Spark within a Data Lakehouse framework, "
        "the proposed system successfully integrates four heterogeneous data sources into a unified, "
        "analysable dataset. The Silver layer\u2019s cleansing operations\u2014industry-specific imputation, "
        "deduplication, derived feature engineering, and NLP contextual mapping\u2014produce a high-quality "
        "dataset suitable for advanced analytics. The Gold layer\u2019s machine learning models deliver "
        "actionable performance profiles and predictions while maintaining strict isolation between "
        "clustering and classification features to prevent data leakage. The design is production-ready, "
        "with automated testing, serialised model artifacts, and comprehensive logging.")

    # Appendix A
    doc.add_page_break()
    add_heading_1(doc, "Appendix A: Project Timeline")
    add_normal_para(doc,
        "The proposed project spans from 20 May 2026 to 10 July 2026, with key milestones "
        "including problem definition and literature review, data collection and pipeline prototyping, "
        "full pipeline implementation with machine learning integration, system testing and validation, "
        "and final report preparation and submission. The Gantt chart below provides a visual "
        "representation of the task schedule and dependencies.",
        first_line_indent=0)

    chart_path = os.path.join(VIZ_DIR, 'gantt_chart.png')
    if os.path.exists(chart_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(chart_path, width=Inches(6.0))

    # Appendix B
    add_heading_1(doc, "Appendix B: Generative AI Use Log")
    g_table = doc.add_table(rows=4, cols=3)
    g_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    gh = g_table.rows[0].cells
    gh[0].text, gh[1].text, gh[2].text = "Date", "Task", "GAI Use / Prompt"
    for c in gh:
        apply_font(c.paragraphs[0], bold=True)

    logs = [
        ("19 June 2026", "Code Refactoring",
         "Used Claude to optimize PySpark joins. Prompt: 'I have a memory bottleneck in my Silver layer join. How can I optimize this in PySpark without Pandas?'"),
        ("20 June 2026", "Proposal Draft",
         "Used ChatGPT to structure the methodology. Prompt: 'Help me outline a methodology section for a CRISP-DM data engineering project based on this rubric.'"),
        ("20 June 2026", "System Design",
         "Used Claude for Star Schema design. Prompt: 'Suggest a Star Schema design where the fact table tracks AI agent executions and dimensions include organization, agent config, and research context.'")
    ]
    for i, (d, t, u) in enumerate(logs):
        r = g_table.rows[i+1].cells
        r[0].text, r[1].text, r[2].text = d, t, u
        for c in r:
            apply_font(c.paragraphs[0])

    # References
    add_heading_1(doc, "References")
    refs = [
        "[1] Armbrust, M., Ghodsi, A., Xin, R., & Zaharia, M. (2021). Lakehouse: A New Generation of Open Platforms that Unify "
         "Data Warehousing and Advanced Analytics. Proceedings of the 11th Conference on Innovative "
         "Data Systems Research (CIDR).",
        "[2] Zaharia, M., Xin, R., Wendell, P., Das, T., Armbrust, M., Dave, A., ... & Stoica, I. (2016). Apache Spark: "
         "A Unified Engine for Big Data Processing. Communications of the ACM, 59(11), 56–65.",
        "[3] Priem, J., Piwowar, H., & Orr, R. (2022). OpenAlex: A Fully Open Index of Scholarly Works, Authors, Venues, "
         "Institutions, and Concepts. Proceedings of the 26th International Conference on Science, Technology and "
         "Innovation Indicators (STI).",
        "[4] Platt, J. C. (1999). Probabilistic Outputs for Support Vector Machines and Comparisons to Regularized "
         "Likelihood Methods. Advances in Large Margin Classifiers, 10(3), 61–74.",
        "[5] Larsen, K. L., & Aone, C. (1999). Fast and Effective Text Mining Using Linear-Time Document Clustering. "
         "Proceedings of the Fifth ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 16–22.",
        "[6] Bolón-Canedo, V., Sánchez-Maroño, N., & Alonso-Betanzos, A. (2013). A Review of Feature Selection "
         "Methods on Synthetic Data. Knowledge and Information Systems, 34(3), 483–519.",
        "[7] Zhang, Y., & Zhang, Z. (2020). Data Lakehouse: A Survey Paper. arXiv preprint arXiv:2010.00874.",
        "[8] Rousseeuw, P. J. (1987). Silhouettes: A Graphical Aid to the Interpretation and Validation of Cluster "
         "Analysis. Journal of Computational and Applied Mathematics, 20, 53–65.",
        "[9] Apache Software Foundation. (2023). Apache Spark MLlib: Scalable Machine Learning Library. "
         "Retrieved from https://spark.apache.org/mllib/",
        "[10] Boyd, S., & Vandenberghe, L. (2004). Convex Optimization. Cambridge University Press."
    ]
    for ref in refs:
        add_normal_para(doc, ref, first_line_indent=0)


def append_id_content(doc):
    make_toc(doc, "Daftar Isi", is_en=False)

    # Abstrak
    doc.add_page_break()
    add_heading_1(doc, "Abstrak")
    add_normal_para(doc,
        "Proyek ini membangun pipeline data untuk mengukur seberapa baik kinerja sistem Agentic AI di perusahaan. "
        "Saat ini, pengambil keputusan tidak punya angka pasti untuk mengevaluasi AI mereka dan hanya mengandalkan "
        "tebakan dari log yang tersebar. Untuk mengatasi ini, kami menggunakan Apache Spark untuk membangun pipeline "
        "berarsitektur Medallion (Bronze, Silver, Gold). Kami menggabungkan 4 jenis data: file CSV berisi 5.500 data "
        "perusahaan, JSON tolok ukur industri, log sistem, dan kutipan OpenAlex. Pipeline ini membersihkan data, mengisi "
        "nilai kosong berdasarkan rata-rata industri, dan mencocokkan jurnal akademik ke sektor bisnis menggunakan NLP. "
        "Di layer Gold, kami memakai K-Means clustering (skor Silhouette 0,5276) untuk mengelompokkan perusahaan, dan "
        "Regresi Logistik (akurasi 56,01%) untuk memprediksi produktivitas. Kami memastikan fitur clustering dipisah dari "
        "model prediksi agar tidak ada kebocoran data (data leakage). Model akhir diekspor dan siap pakai, memberikan "
        "angka nyata bagi pimpinan alih-alih sekadar asumsi.")

    # Bab 1
    doc.add_page_break()
    add_heading_1(doc, "1. Pendahuluan")

    add_normal_para(doc,
        "Agentic AI mulai mengubah cara perusahaan bekerja. Berbeda dengan bot biasa yang cuma mengikuti "
        "skrip, sistem ini bisa merencanakan, menggunakan alat, dan beradaptasi sendiri. Perusahaan sekarang "
        "memakai Agentic AI untuk layanan pelanggan, rantai pasok, dan operasi TI\u2014terutama untuk "
        "menekan biaya dan mempercepat keputusan.")

    add_normal_para(doc,
        "Tapi masalahnya: sangat sulit membuktikan bahwa agen AI ini benar-benar efektif. Manajer butuh "
        "angka pasti, tapi datanya sering berserakan. Log sistem tersimpan di file teks, metrik perusahaan "
        "di CSV, dan standar industri ada di API. Tanpa satu pipeline terpusat untuk membersihkan dan "
        "menggabungkan data ini, perusahaan pada akhirnya hanya menebak apakah AI mereka benar-benar bekerja.")

    add_normal_para(doc,
        "Proyek ini menyelesaikan masalah pengukuran tersebut. Kami membangun pipeline Medallion "
        "berbasis Spark untuk mengubah log mentah menjadi profil kinerja yang jelas. Lapisan Bronze "
        "menyimpan data mentah. Lapisan Silver membersihkannya, mengisi nilai kosong, dan "
        "menggunakan NLP untuk memetakan kutipan akademik ke sektor bisnis. Terakhir, lapisan "
        "Gold menerapkan machine learning untuk memprediksi kinerja. Hasil akhirnya adalah dataset "
        "andal yang membantu pimpinan membuat keputusan berbasis data tentang sistem AI mereka.")

    add_heading_2(doc, "1.1 Domain Bisnis & Identifikasi Masalah")
    add_normal_para(doc,
        "Domain bisnis yang dipilih untuk proyek ini adalah Professional Services, sebuah sektor "
        "yang mencakup firma konsultasi, praktik hukum, penyedia layanan TI, dan organisasi "
        "konsultasi keuangan. Perusahaan-perusahaan ini dicirikan oleh operasi padat pengetahuan "
        "di mana keahlian manusia dilengkapi dengan sistem teknologi yang semakin otonom. Dalam "
        "domain ini, aplikasi Agentic AI diterapkan untuk tugas-tugas seperti tinjauan dokumen "
        "otomatis, analisis kontrak, triase komunikasi klien, pembuatan rekomendasi berbasis data, "
        "dan otomatisasi proses cerdas. Kinerja sistem ini secara langsung berdampak pada kepuasan "
        "klien, efisiensi operasional, dan kepatuhan regulasi, sehingga kerangka evaluasi yang "
        "andal menjadi kebutuhan organisasi yang kritikal.")

    add_normal_para(doc,
        "Meskipun ada kebutuhan yang jelas untuk evaluasi kinerja, praktik saat ini mengungkapkan "
        "masalah mendasar: tidak ada kerangka rekayasa data yang terstandarisasi untuk menyatukan "
        "telemetri terfragmentasi yang dihasilkan oleh sistem Agentic AI. Data operasional "
        "diproduksi dalam berbagai format yang tidak terhubung\u2014ekspor CSV terstruktur dari "
        "dasbor penerapan, respons JSON semi-terstruktur dari API benchmarking eksternal, file "
        "log aplikasi mentah yang mencatat metrik eksekusi tingkat sistem, dan data kutipan "
        "akademik dari database penelitian publik. Masing-masing sumber ini menggunakan skema, "
        "konvensi penamaan, dan profil kualitas data yang berbeda. Tanpa integrasi sistematis, "
        "tim kepemimpinan tidak dapat melakukan analisis lintas-sumber, mengagregasi tren kinerja, "
        "atau menghubungkan efisiensi operasional dengan standar eksternal. Hal ini sangat "
        "membatasi kemampuan organisasi untuk membuat keputusan yang terinformasi dan berbasis "
        "bukti tentang strategi dan investasi AI mereka.")

    add_heading_2(doc, "1.2 Tujuan Proyek")
    add_normal_para(doc,
        "Tujuan utama dari proyek ini adalah untuk merancang dan mengimplementasikan pipeline "
        "rekayasa data otomatis dan terukur menggunakan Apache Spark yang mengintegrasikan data "
        "telemetri multi-sumber untuk menghasilkan profil kinerja kepemimpinan Agentic AI yang "
        "andal. Tujuan khususnya adalah sebagai berikut:",
        first_line_indent=0)

    objs = [
        "1. Membangun pipeline data berstruktur Medallion (Bronze, Silver, Gold) menggunakan Apache "
           "Spark yang mengumpulkan, menstandardisasi, dan mengintegrasikan empat sumber data "
           "heterogen\u2014dataset CSV utama berisi 5.500 catatan organisasi, file JSON tolok ukur "
           "industri eksternal, log eksekusi sistem dalam format CSV, dan data kutipan akademik "
           "dari API OpenAlex.",

        "2. Menyusun pipeline pembersihan data otomatis yang menggunakan Spark User-Defined "
           "Functions dan operasi agregasi berbasis window untuk melakukan imputasi median spesifik "
           "industri untuk nilai null, deduplikasi, standarisasi tipe, rekayasa fitur turunan, "
           "dan pemetaan kontekstual berbasis NLP dari kutipan makalah ilmiah ke sektor bisnis "
           "yang cocok.",

        "3. Menerapkan dan memvalidasi model klasterisasi K-Means tanpa pengawasan dan classifier "
           "Regresi Logistik multikelas yang di-tune di lapisan pemrosesan Gold, memastikan "
           "isolasi data antara fitur klasterisasi dan input klasifikasi untuk mencegah data "
           "leakage, dan mengekspor model berkinerja terbaik sebagai artefak serialisasi siap produksi."
    ]
    for obj in objs:
        add_normal_para(doc, obj)

    add_heading_2(doc, "1.3 Analisis Risiko & Rencana Mitigasi")
    add_normal_para(doc,
        "Implementasi pipeline rekayasa data untuk evaluasi Agentic AI membawa beberapa risiko yang melekat. "
        "Mengidentifikasi risiko-risiko ini secara dini memastikan bahwa strategi mitigasi yang tepat "
        "dapat diintegrasikan ke dalam arsitektur sistem dan jadwal proyek.")
        
    r_table = doc.add_table(rows=4, cols=3)
    r_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_hdr = r_table.rows[0].cells
    r_hdr[0].text, r_hdr[1].text, r_hdr[2].text = "Risiko", "Dampak / Probabilitas", "Strategi Mitigasi"
    for c in r_hdr:
        apply_font(c.paragraphs[0], bold=True)

    risks_id = [
        ("Kualitas Data & Pergeseran Skema", "Dampak: Tinggi\nProbabilitas: Sedang", 
         "Menerapkan imputasi median untuk nilai numerik yang hilang dan memvalidasi integritas skema menggunakan unit test PySpark otomatis sebelum dimuat ke lapisan Gold."),
        ("Keterbatasan Alat & Perangkat Keras", "Dampak: Tinggi\nProbabilitas: Tinggi", 
         "Lingkungan memiliki RAM terbatas (1.9GB. Dimitigasi dengan mengatur jumlah partition Spark, menghindari Python UDF untuk mencegah overhead serialisasi memori, dan menulis dalam chunk batch."),
        ("Alikan Pemangku Kepentingan", "Dampak: Sedang\nProbabilitas: Rendah", 
         "Memastikan output ML sejajar dengan KPI bisnis dengan memvalidasi hasil klasterisasi tanpa pengawasan (Skor Silhouette) terhadap metrik kepercayaan kepemimpinan.")
    ]
    for i, (r_name, r_imp, r_mit) in enumerate(risks_id):
        row = r_table.rows[i+1].cells
        row[0].text, row[1].text, row[2].text = r_name, r_imp, r_mit
        for c in row:
            apply_font(c.paragraphs[0])
            
    doc.add_paragraph()

    # Bab 2
    doc.add_page_break()
    add_heading_1(doc, "2. Tinjauan Pustaka")
    add_normal_para(doc,
        "Pipeline data yang solid sangat penting untuk menangani volume data perusahaan yang besar. "
        "Data Warehouse tradisional terlalu kaku untuk log semi-terstruktur, sementara Data Lake kurang "
        "aman untuk transaksi. Arsitektur Data Lakehouse mengatasi masalah ini dengan menggabungkan "
        "fleksibilitas lake dan keandalan warehouse. Ia menggunakan format kolom seperti Parquet dan mesin "
        "komputasi seperti Apache Spark untuk memproses data dan menjalankan machine learning secara langsung (Armbrust et al., 2021).")

    add_normal_para(doc,
        "Proyek ini dibangun berdasarkan tiga temuan utama. Pertama, model Lakehouse mengurangi duplikasi data "
        "dan menyederhanakan pipeline (Armbrust et al., 2021). Kedua, pemrosesan in-memory Apache Spark sangat "
        "ideal untuk menggabungkan dataset besar yang kompleks (Zaharia et al., 2016)\u2014tepat seperti yang kami "
        "butuhkan untuk log Agentic AI. Terakhir, riset terbaru menyoroti bahwa menggabungkan log sistem dengan "
        "data eksternal\u2014seperti kutipan OpenAlex\u2014sangat penting untuk mengevaluasi keandalan agen AI di "
        "dunia nyata (Priem et al., 2022). Tiga paper ini yang mendasari pilihan arsitektur kami.")

    # Bab 3
    doc.add_page_break()
    add_heading_1(doc, "3. Metodologi")

    add_normal_para(doc,
        "Bagian ini menguraikan cara kami merancang, membangun, dan menguji pipeline data. "
        "Metodologi yang jelas memastikan proyek mencapai tujuannya dan mudah direproduksi nanti. "
        "Kami membahas pilihan framework, empat sumber data yang dipakai, proses ETL, keputusan "
        "penyimpanan, dan implementasi machine learning di layer Gold.")

    add_normal_para(doc,
        "Kami menggunakan pola Arsitektur Medallion. Lapisan Bronze menyimpan data mentah apa adanya. "
        "Lapisan Silver membersihkan dan mengintegrasikan data tersebut. Lapisan Gold menerapkan "
        "analitik dan machine learning. Pengaturan ini menjaga kode tetap rapi, menangani data "
        "buruk secara bertahap, dan sudah jadi standar industri. Keseluruhan proyek ini kami "
        "kelola menggunakan kerangka kerja CRISP-DM.")

    add_heading_2(doc, "3.1 Pemilihan Metodologi & Justifikasi")
    add_normal_para(doc,
        "Proyek ini mengadopsi Cross-Industry Standard Process for Data Mining (CRISP-DM) "
        "sebagai kerangka metodologis utamanya. CRISP-DM dipilih karena merupakan kerangka kerja "
        "yang paling banyak digunakan dan matang untuk proyek data mining dan analitik, "
        "menyediakan siklus hidup yang terstruktur namun fleksibel yang dapat mengakomodasi "
        "komponen rekayasa data dan pembelajaran mesin. Kerangka kerja ini terdiri dari enam "
        "fase iteratif: Pemahaman Bisnis, Pemahaman Data, Persiapan Data, Pemodelan, Evaluasi, "
        "dan Penerapan. Untuk proyek ini, fase Pemahaman Bisnis sesuai dengan formulasi masalah "
        "yang dijelaskan di Bagian 1, fase Pemahaman Data dan Persiapan Data sesuai dengan "
        "langkah pemrosesan lapisan Bronze dan Silver, dan fase Pemodelan dan Evaluasi sesuai "
        "dengan analitik lapisan Gold. Sifat iteratif dari CRISP-DM sangat berharga karena "
        "memungkinkan tahap persiapan data dan pemodelan saling memberi informasi\u2014kemampuan "
        "penting saat mengintegrasikan sumber data yang beragam dengan skema dan profil kualitas "
        "yang berbeda.")

    add_heading_2(doc, "3.2 Identifikasi Sumber Data")
    add_normal_para(doc,
        "Proyek ini menggunakan empat sumber data yang berbeda untuk menangkap konteks operasional "
        "penuh dari penerapan Agentic AI. Sumber-sumber ini dipilih untuk mewakili tiga kategori "
        "utama data yang biasanya ditemui organisasi: catatan operasional terstruktur internal (S1), "
        "standar benchmarking eksternal (S2), telemetri eksekusi sistem granular (S3), dan konteks "
        "penelitian akademik (S4). Tabel berikut merangkum setiap sumber, formatnya, dan perannya "
        "dalam pipeline. Deskripsi tingkat kolom yang terperinci disediakan di paragraf berikutnya "
        "untuk mendukung pembaca dalam memahami skema yang mengalir melalui setiap lapisan Medallion.")

    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'ID', 'Sumber Data', 'Format', 'Deskripsi'
    for c in hdr:
        apply_font(c.paragraphs[0], bold=True)

    sources = [
        ("S1", "Dataset Kepemimpinan Agentic AI", "CSV",
         "Dataset operasional utama berisi 5.500 catatan metrik organisasi, tingkat otonomi, tingkat keberhasilan tugas, dan skor kepercayaan kepemimpinan dari penerapan Agentic AI di berbagai industri."),
        ("S2", "Tolok Ukur Kepemimpinan Eksternal", "JSON",
         "Tolok ukur dasar industri yang menyediakan nilai target referensi untuk tingkat keberhasilan tugas dan ambang produktivitas."),
        ("S3", "Log Eksekusi Agen", "CSV",
         "Log eksekusi sistem tingkat granular yang mencatat persentase penggunaan CPU, konsumsi memori dalam MB, waktu eksekusi, dan jumlah kesalahan aplikasi per catatan organisasi."),
        ("S4", "Publikasi Akademik OpenAlex", "API JSON",
         "Metadata akademik dari database akses terbuka OpenAlex berisi 25 catatan publikasi ilmiah dengan jumlah kutipan, tahun publikasi, dan kata kunci judul.")
    ]
    for i, s in enumerate(sources):
        r = table.rows[i+1].cells
        r[0].text, r[1].text, r[2].text, r[3].text = s[0], s[1], s[2], s[3]
        for c in r:
            apply_font(c.paragraphs[0])

    doc.add_paragraph()

    add_normal_para(doc,
        "Sumber S1, Dataset Kepemimpinan Agentic AI, adalah input utama ke pipeline. "
        "Dataset ini disimpan dalam format CSV (Comma-Separated Values) dan berisi 5.500 "
        "catatan, masing-masing mewakili unit organisasi yang telah menerapkan sistem Agentic AI. "
        "Dataset ini mencakup 15 kolom yang menangkap konteks organisasi, karakteristik agen, "
        "dan metrik kinerja operasional. Kolom konteks organisasi mencakup Record_ID (pengenal "
        "alfanumerik unik untuk setiap catatan), Organization_Name (nama perusahaan yang "
        "menerapkan sistem AI), Industry (klasifikasi sektor bisnis seperti Keuangan, Kesehatan, "
        "atau Pendidikan), Organization_Size (dikategorikan sebagai Kecil, Menengah, atau Besar), "
        "dan AI_Maturity_Level (variabel kategorikal yang menunjukkan tingkat pengalaman "
        "organisasi dengan AI).")

    add_normal_para(doc,
        "Kolom konfigurasi agen menggambarkan karakteristik sistem AI yang diterapkan dan "
        "mencakup Agent_Type (jenis kerangka kerja seperti LangChain, AutoGPT, atau CrewAI), "
        "Use_Case_Area (domain fungsional seperti Dukungan Pelanggan, Analisis Data, atau "
        "Generasi Kode), Agent_Autonomy_Level (peringkat kategorikal kemandirian operasional "
        "agen), Decision_Making_Type (pendekatan algoritmik untuk pengambilan keputusan seperti "
        "Berbasis Aturan, Berbasis LLM, atau Hibrida), dan Context_Awareness_Score (skor numerik "
        "antara 0 dan 100 yang mengukur kemampuan agen untuk memahami konteks operasionalnya). "
        "Kolom metrik kinerja mencakup Task_Success_Rate (persentase tugas yang diberikan "
        "diselesaikan dengan sukses), Productivity_Improvement_Percent (peningkatan output "
        "operasional sejak penerapan), Response_Time_Seconds (waktu respons rata-rata agen), "
        "Complexity_Score (ukuran normal dari kesulitan tugas), dan Leadership_Trust_Score "
        "(indeks komposit yang mencerminkan kepercayaan eksekutif terhadap keandalan sistem). "
        "Kelima belas kolom ini secara kolektif memberikan pandangan multi-dimensi yang kaya "
        "tentang pengalaman setiap organisasi dengan Agentic AI.")

    add_heading_2(doc, "3.3 Ingesti Data, Pembersihan & Pemrosesan Dalam Memori")
    add_normal_para(doc,
        "Proses ingesti data dimulai dengan lapisan Bronze, di mana empat file sumber dimuat "
        "ke dalam Spark DataFrames menggunakan pembaca khusus format. File CSV dibaca dengan "
        "inferensi header dan penegakan skema yang diaktifkan, file JSON diuraikan dengan "
        "dukungan multi-baris untuk struktur bersarang, dan respons API OpenAlex dikonsumsi "
        "sebagai file JSON standar. Setiap dataset yang diingesti diperkaya dengan bidang "
        "metadata yang mencatat stempel waktu ingesti dan pengenal sumber sebelum ditulis ke "
        "file Parquet terpartisi di direktori output Bronze. Lapisan penyimpanan mentah ini "
        "berfungsi sebagai arsip permanen yang tidak dimodifikasi yang memungkinkan pelacakan "
        "garis keturunan data dan pemrosesan ulang jika diperlukan.")

    add_normal_para(doc,
        "Di lapisan Silver, operasi pembersihan data diterapkan secara sistematis. Pertama, "
        "catatan dideduplikasi dengan mengidentifikasi dan menghapus entri duplikat berdasarkan "
        "bidang Record_ID. Kedua, semua kolom numerik dikonversi ke tipe data Spark yang sesuai "
        "untuk memastikan konsistensi komputasional. Ketiga, bidang teks kategorikal "
        "distandarisasi ke huruf kecil untuk menghilangkan masalah sensitivitas huruf selama "
        "operasi join. Keempat, nilai numerik yang hilang diimputasi menggunakan nilai median "
        "yang dihitung untuk setiap kelompok industri; pendekatan ini mempertahankan karakteristik "
        "distribusional spesifik industri daripada menerapkan imputasi global yang dapat membiaskan "
        "hasil. Data log eksekusi digabungkan ke dataset utama pada Record_ID, memungkinkan "
        "derivasi fitur efisiensi seperti Memory_Per_Message_MB. Akhirnya, judul makalah OpenAlex "
        "diproses melalui skrip klasifikasi berbasis regex yang memetakan setiap publikasi ke "
        "sektor bisnis yang cocok.")

    add_heading_2(doc, "3.4 Penyimpanan Data & Justifikasi")
    add_normal_para(doc,
        "Semua data yang diproses disimpan dalam format kolom Apache Parquet di tiga lapisan "
        "direktori Medallion: Bronze (arsip mentah), Silver (data yang dibersihkan), dan Gold "
        "(output analitis). Parquet dipilih sebagai format penyimpanan karena beberapa alasan "
        "teknis. Tata letak penyimpanan kolomnya memungkinkan predicate pushdown dan column "
        "projection selama kueri Spark SQL, yang secara signifikan mengurangi I/O dengan hanya "
        "membaca kolom yang diperlukan untuk operasi tertentu. Pengkodean kamus lebih lanjut "
        "mengompresi nilai string berulang, mengurangi jejak penyimpanan. Penegakan skema "
        "memastikan bahwa data yang ditulis ke setiap lapisan sesuai dengan struktur yang "
        "diharapkan, mencegah korupsi diam-diam dari penyimpangan skema.")

    add_heading_2(doc, "3.5 Integrasi Algoritma Kecerdasan Buatan (AI)")
    add_normal_para(doc,
        "Untuk memberikan nilai analitis di luar agregasi dasar, lapisan Gold mengintegrasikan "
        "dua model pembelajaran mesin dari pustaka PySpark MLlib. Model-model ini beroperasi "
        "pada dataset yang telah dibersihkan dan terpadu yang dihasilkan oleh lapisan Silver "
        "dan dirancang untuk menghasilkan wawasan yang dapat ditindaklanjuti bagi pengambil "
        "keputusan bisnis.")

    add_normal_para(doc,
        "Model pertama adalah algoritma klasterisasi K-Means tanpa pengawasan yang dikonfigurasi "
        "untuk membagi 5.500 catatan organisasi ke dalam tiga tingkatan kinerja yang berbeda: "
        "High Performer, Average Performer, dan Low Performer. Klasterisasi dilakukan pada tiga "
        "fitur yang dinormalisasi: Task_Success_Rate, Productivity_Improvement_Percent, dan "
        "Leadership_Trust_Score. Fitur-fitur ini pertama-tama dirakit menjadi feature vector "
        "menggunakan VectorAssembler, kemudian distandarisasi menggunakan StandardScaler untuk "
        "memastikan bobot yang sama selama komputasi jarak. Algoritma K-Means dikonfigurasi "
        "dengan k=3 klaster dan maksimum 20 iterasi. Kualitas model dievaluasi menggunakan "
        "skor Silhouette, yang mengukur seberapa mirip setiap titik dengan klusternya sendiri "
        "dibandingkan dengan klaster lain. Skor Silhouette yang dicapai adalah 0,5276.")

    add_normal_para(doc,
        "Model kedua adalah classifier Regresi Logistik multikelas terawasi yang memprediksi "
        "Productivity_Category (Tinggi, Rata-rata, atau Rendah) berdasarkan fitur operasional "
        "tingkat sistem: Context_Awareness_Score, Response_Time_Seconds, Task_Complexity_Score, "
        "Memory_Per_Message_MB, dan CPU_Utilization_Percent. Fitur-fitur ini sengaja dipilih "
        "untuk mengecualikan metrik klasterisasi, sehingga mencegah data leakage. Model di-tune "
        "menggunakan CrossValidator dengan 3-fold cross-validation. Model yang di-tune mencapai "
        "akurasi 56,01%, F1-score 50,29%, presisi 52,08%, dan recall 56,01%. Pipeline model "
        "terbaik, termasuk bobot yang dilatih dan parameter scaler, diserialisasi dan diekspor "
        "ke file fisik di direktori output Gold.")

    # Bab 4
    doc.add_page_break()
    add_heading_1(doc, "4. Desain Sistem & Arsitektur")

    add_normal_para(doc,
        "Arsitektur sistem untuk proyek ini dirancang untuk menangani siklus hidup ujung-ke-ujung dari "
        "data telemetri Agentic AI. Dari sumber data yang mentah dan terfragmentasi menjadi skema "
        "multidimensi yang terstruktur membutuhkan desain yang tangguh dan toleran terhadap kegagalan. "
        "Apache Spark menjadi mesin komputasi inti, dipilih secara spesifik karena kemampuan pemrosesan "
        "in-memory terdistribusinya yang menghilangkan bottleneck Out-Of-Memory yang biasa terjadi pada "
        "alat single-node seperti Pandas saat melakukan join antara log eksekusi besar dengan respons "
        "JSON API eksternal.")

    add_normal_para(doc,
        "Bab ini merinci struktur desain pipeline, terbagi menjadi lima komponen inti. Dimulai dari "
        "diagram Use Case yang menggambarkan interaksi pengguna, dilanjutkan dengan alur data ETL "
        "dalam paradigma Medallion Lakehouse. Sub-bab selanjutnya merinci desain database logis "
        "Star Schema, antarmuka dashboard analitis yang ditujukan untuk konsumsi kepemimpinan, "
        "serta strategi pengujian unit yang ketat yang digunakan untuk menjamin zero data leakage "
        "antara fase klasterisasi K-Means dan klasifikasi Regresi Logistik.")

    add_heading_2(doc, "4.1 Deskripsi Use Case Diagram")
    add_normal_para(doc,
        "Use Case Sistem melibatkan dua aktor utama: Data Engineer dan Data Analyst. "
        "Data Engineer bertanggung jawab untuk memulai eksekusi pipeline, memantau status "
        "pemrosesan di ketiga lapisan Medallion, memeriksa file log untuk kesalahan, dan "
        "memicu pemrosesan ulang jika masalah kualitas data terdeteksi. Data Analyst "
        "berinteraksi dengan output lapisan Gold untuk melakukan kueri intelijen bisnis, "
        "memeriksa profil klasterisasi, mengevaluasi prediksi klasifikasi, dan menghasilkan "
        "dasbor visualisasi untuk pelaporan kepemimpinan. Model pembelajaran mesin dieksekusi "
        "secara otomatis sebagai bagian dari orkestrasi pipeline.")
        
    if os.path.exists(CHART_UC):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(CHART_UC, width=Inches(5.0))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.add_run("Gambar 4.1.1: Use Case Diagram Sistem Lakehouse").italic = True

    add_heading_2(doc, "4.2 Arsitektur Sistem (Aliran ETL & Penyimpanan)")
    add_normal_para(doc,
        "Arsitektur pipeline mengikuti desain Medallion modular. Dataset mentah disimpan "
        "dalam struktur folder lokal dan diproses secara berurutan oleh Apache Spark. Lapisan "
        "Bronze menghasilkan arsip Parquet yang diaudit, lapisan Silver menghasilkan output "
        "Parquet yang telah dibersihkan dan direkayasa fitur, dan lapisan Gold menghasilkan "
        "mart analitis yang diagregasi bersama dengan prediksi pembelajaran mesin dan artefak "
        "model yang diserialisasi. Setiap lapisan beroperasi sebagai job Spark independen "
        "dengan validasi input, logika transformasi, dan langkah verifikasi outputnya sendiri. "
        "Pencatatan log dilakukan di setiap tahap untuk memberikan visibilitas operasional.")

    add_heading_2(doc, "4.3 Desain Database & Skema Multidimensi (Star Schema)")
    add_normal_para(doc,
        "Untuk pemodelan analitis multidimensi, Lapisan Gold disusun sebagai database Star "
        "Schema. Tabel fakta pusat, Fact_AI_Deployment, berisi referensi kunci asing ke tiga "
        "tabel dimensi: Dim_Organization (atribut organisasi), Dim_Agent (detail konfigurasi "
        "agen AI), dan Dim_Research_Context (metrik kutipan ilmiah hasil pemetaan NLP). Desain "
        "dimensional ini memungkinkan analis bisnis untuk mengkueri metrik kinerja di beberapa "
        "sumbu analitis tanpa memerlukan join multi-tabel yang kompleks.")
        
    if os.path.exists(CHART_ERD):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(CHART_ERD, width=Inches(5.5))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.add_run("Gambar 4.3.1: Diagram ERD Star Schema").italic = True

    add_heading_2(doc, "4.4 Antarmuka Utama & Desain Dashboard Analitis")
    add_normal_para(doc,
        "Antarmuka analitis dirancang sebagai tiga dasbor visualisasi. Dasbor pertama "
        "menampilkan bagan perbandingan tingkat industri yang memplot Skor Kepercayaan "
        "Kepemimpinan terhadap Peningkatan Produktivitas di semua sektor. Dasbor kedua "
        "menampilkan distribusi profil klasterisasi. Dasbor ketiga menyajikan metrik akurasi "
        "prediksi model klasifikasi. Dasbor ini dimaksudkan untuk diimplementasikan dalam "
        "alat intelijen bisnis seperti Power BI atau Tableau.")

    add_heading_2(doc, "4.5 Strategi Pengujian & Validasi")
    add_normal_para(doc,
        "Untuk memastikan kualitas data dan keandalan pipeline, test suite otomatis dieksekusi "
        "setelah setiap proses pipeline. Suite ini memverifikasi empat aspek: keselarasan skema "
        "(memeriksa bahwa semua kolom turunan yang diharapkan ada dengan tipe data yang benar), "
        "pernyataan nilai null (memverifikasi bahwa bidang kritis berisi nol nilai null setelah "
        "imputasi lapisan Silver), integritas catatan (mengkonfirmasi bahwa tepat 5.500 catatan "
        "dipertahankan dari Bronze hingga Gold), dan validasi artefak model (memeriksa bahwa "
        "file model yang diserialisasi ada di jalur output yang diharapkan).")

    # Bab 5
    doc.add_page_break()
    add_heading_1(doc, "5. Kesimpulan")
    add_normal_para(doc,
        "Proposal ini telah menyajikan desain komprehensif untuk pipeline rekayasa data yang "
        "terukur dan otomatis yang mengatasi kesenjangan kritikal dalam evaluasi kinerja "
        "Agentic AI. Dengan mengimplementasikan arsitektur Medallion pada Apache Spark dalam "
        "kerangka Data Lakehouse, sistem yang diusulkan berhasil mengintegrasikan empat sumber "
        "data heterogen ke dalam dataset terpadu yang dapat dianalisis. Desain siap produksi, "
        "dengan pengujian otomatis, artefak model yang diserialisasi, dan pencatatan log "
        "yang komprehensif.")

    # Lampiran A
    doc.add_page_break()
    add_heading_1(doc, "Lampiran A: Timeline Proyek")
    add_normal_para(doc,
        "Proyek yang diusulkan berlangsung dari 20 Mei 2026 hingga 10 Juli 2026, dengan "
        "pencapaian utama termasuk definisi masalah dan tinjauan pustaka, pengumpulan data "
        "dan pembuatan prototipe pipeline, implementasi pipeline penuh dengan integrasi "
        "pembelajaran mesin, pengujian dan validasi sistem, serta persiapan laporan akhir "
        "dan pengumpulan. Bagan Gantt di bawah ini menyediakan representasi visual dari "
        "jadwal tugas dan dependensi.",
        first_line_indent=0)

    chart_path = os.path.join(VIZ_DIR, 'gantt_chart.png')
    if os.path.exists(chart_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(chart_path, width=Inches(6.0))

    # Lampiran B
    add_heading_1(doc, "Lampiran B: Log Penggunaan AI Generatif")
    g_table = doc.add_table(rows=4, cols=3)
    g_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    gh = g_table.rows[0].cells
    gh[0].text, gh[1].text, gh[2].text = "Tanggal", "Tugas", "Penggunaan GAI / Prompt"
    for c in gh:
        apply_font(c.paragraphs[0], bold=True)

    logs = [
        ("19 Juni 2026", "Code Refactoring",
         "Pakai Claude untuk optimasi strategi join PySpark. Prompt: 'Saya ada bottleneck memori di proses join Silver layer. Gimana cara optimasinya di PySpark tanpa pakai Pandas?'"),
        ("20 Juni 2026", "Draf Proposal",
         "Pakai ChatGPT untuk menstrukturisasi bagian metodologi. Prompt: 'Tolong buatkan outline bagian metodologi untuk proyek data engineering CRISP-DM berdasarkan rubrik ini.'"),
        ("20 Juni 2026", "Desain Sistem",
         "Pakai Claude untuk brainstorming struktur Star Schema. Prompt: 'Beri saran desain Star Schema di mana tabel fakta melacak eksekusi agen AI dan dimensinya mencakup organisasi, config agen, dan konteks riset.'")
    ]
    for i, (d, t, u) in enumerate(logs):
        r = g_table.rows[i+1].cells
        r[0].text, r[1].text, r[2].text = d, t, u
        for c in r:
            apply_font(c.paragraphs[0])

    # Daftar Pustaka
    add_heading_1(doc, "Daftar Pustaka")
    refs = [
        "[1] Armbrust, M., Ghodsi, A., Xin, R., & Zaharia, M. (2021). Lakehouse: A New Generation of Open Platforms that Unify "
         "Data Warehousing and Advanced Analytics. Proceedings of the 11th Conference on Innovative "
         "Data Systems Research (CIDR).",
        "[2] Zaharia, M., Xin, R., Wendell, P., Das, T., Armbrust, M., Dave, A., ... & Stoica, I. (2016). Apache Spark: "
         "A Unified Engine for Big Data Processing. Communications of the ACM, 59(11), 56–65.",
        "[3] Priem, J., Piwowar, H., & Orr, R. (2022). OpenAlex: A Fully Open Index of Scholarly Works, Authors, Venues, "
         "Institutions, and Concepts. Proceedings of the 26th International Conference on Science, Technology and "
         "Innovation Indicators (STI).",
        "[4] Platt, J. C. (1999). Probabilistic Outputs for Support Vector Machines and Comparisons to Regularized "
         "Likelihood Methods. Advances in Large Margin Classifiers, 10(3), 61–74.",
        "[5] Larsen, K. L., & Aone, C. (1999). Fast and Effective Text Mining Using Linear-Time Document Clustering. "
         "Proceedings of the Fifth ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 16–22.",
        "[6] Bolón-Canedo, V., Sánchez-Maroño, N., & Alonso-Betanzos, A. (2013). A Review of Feature Selection "
         "Methods on Synthetic Data. Knowledge and Information Systems, 34(3), 483–519.",
        "[7] Zhang, Y., & Zhang, Z. (2020). Data Lakehouse: A Survey Paper. arXiv preprint arXiv:2010.00874.",
        "[8] Rousseeuw, P. J. (1987). Silhouettes: A Graphical Aid to the Interpretation and Validation of Cluster "
         "Analysis. Journal of Computational and Applied Mathematics, 20, 53–65.",
        "[9] Apache Software Foundation. (2023). Apache Spark MLlib: Scalable Machine Learning Library. "
         "Retrieved from https://spark.apache.org/mllib/",
        "[10] Boyd, S., & Vandenberghe, L. (2004). Convex Optimization. Cambridge University Press."
    ]
    for ref in refs:
        add_normal_para(doc, ref, first_line_indent=0)


def generate():
    # English
    en_doc = Document(TEMPLATE_PATH)
    append_en_content(en_doc)
    en_path = os.path.join(REPORT_DIR, "Richard_Clay_EN.docx")
    en_doc.save(en_path)
    print(f"EN Proposal saved: {en_path}")

    # Indonesian
    id_doc = Document(TEMPLATE_PATH)
    append_id_content(id_doc)
    id_path = os.path.join(REPORT_DIR, "Richard_Clay_ID.docx")
    id_doc.save(id_path)
    print(f"ID Proposal saved: {id_path}")


if __name__ == "__main__":
    generate()
