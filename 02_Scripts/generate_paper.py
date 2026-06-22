# -*- coding: utf-8 -*-
"""
generate_paper.py — Proceeding Paper Generator (EN + ID) v2
Menggunakan template sebagai fondasi, replace placeholder in-place,
lalu append konten baru di section 2-column terakhir.
"""
import os, shutil, copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
REPORT_DIR = os.path.abspath(os.path.join(BASE_DIR, '..', '07_Reports'))
VIZ_DIR    = os.path.abspath(os.path.join(BASE_DIR, '..', '06_Visualization'))
TEMPLATE   = r"C:\Users\richa\Documents\Sem.6\DatEng_Individual_Project\Materials\Proceeding Paper (Template).docx"
OUT_EN     = os.path.join(REPORT_DIR, "Richard_Clay_Paper_EN.docx")
OUT_ID     = os.path.join(REPORT_DIR, "Richard_Clay_Paper_ID.docx")

UC_PATH   = os.path.join(VIZ_DIR, 'use_case_diagram.png')
ERD_PATH  = os.path.join(VIZ_DIR, 'star_schema_erd.png')
GANTT_PATH= os.path.join(VIZ_DIR, 'gantt_chart.png')
ETL_PATH  = os.path.join(VIZ_DIR, 'etl_architecture.png')

BLACK = RGBColor(0,0,0)

def set_run(r, sz=10, bold=False, italic=False):
    r.font.name = 'Times New Roman'
    r.font.size = Pt(sz)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = BLACK

def fmt_p(p, a=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3, before=0, ls=1.0, indent=0.0):
    p.alignment = a
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = ls
    if indent > 0:
        p.paragraph_format.first_line_indent = Inches(indent)

def add_body(doc, text, sz=10, after=3, indent=0.2):
    p = doc.add_paragraph(style='Body Text')
    p.clear()
    fmt_p(p, after=after, indent=indent)
    r = p.add_run(text)
    set_run(r, sz)

def add_heading1(doc, text, sz=10):
    p = doc.add_paragraph(style='Heading 1')
    p.clear()
    fmt_p(p, a=WD_ALIGN_PARAGRAPH.LEFT, after=3, before=6, indent=0)
    r = p.add_run(text)
    set_run(r, sz, bold=True)

def add_heading2(doc, text, sz=10):
    p = doc.add_paragraph(style='Heading 2')
    p.clear()
    fmt_p(p, a=WD_ALIGN_PARAGRAPH.LEFT, after=2, before=4, indent=0)
    r = p.add_run(text)
    set_run(r, sz, italic=True)

def add_fig_caption(doc, text):
    p = doc.add_paragraph(style='figure caption')
    p.clear()
    fmt_p(p, a=WD_ALIGN_PARAGRAPH.CENTER, after=4, before=2, indent=0)
    r = p.add_run(text)
    set_run(r, 8)

def add_img(doc, path, width_inches=3.1):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt_p(p, after=0, before=0)
        r = p.add_run()
        r.add_picture(path, width=Inches(width_inches))

# ─── Content ──────────────────────────────────────────────────────

PAPER_TITLE = "A Data Engineering Pipeline for Trust Assessment of Agentic AI in Professional Services"
AUTHORS = "Richard Clay, Jonathan H. P. Simatupang, Bens Pardamean"
AFFILIATION = "Computer Science Department, BINUS University, Jakarta, Indonesia 11480"
EMAIL = "richard.clay@binus.ac.id, {jonathan.simatupang, bpardamean}@binus.edu"

EN_ABSTRACT = (
    "The rapid adoption of autonomous AI agents in professional services—law, consulting, accounting—"
    "has created an urgent need for structured trust assessment frameworks. Unlike traditional software, "
    "Agentic AI systems operate with significant autonomy: they triage client emails, monitor infrastructure, "
    "support financial reviews, and coordinate dynamic workflows. Without systematic monitoring, firms risk "
    "deploying agents whose outputs are unreliable or misaligned with organisational standards. This paper "
    "addresses that gap by building a data engineering pipeline that integrates operational telemetry, "
    "benchmark performance data, and external academic citation metrics into a single analytical model. "
    "The pipeline follows a Medallion architecture (Bronze-Silver-Gold) built on Apache Spark, processing "
    "four heterogeneous sources: structured CSV files from Kaggle (2,500 records), JSON API benchmarks from "
    "SeedBench (1,500 records), unstructured agent execution logs (2,000 records), and external OpenAlex "
    "citation metadata via REST API (1,000 records). The Silver layer applies rule-based NLP contextual "
    "mapping to enrich research context tokens using keyword heuristics, while the Gold layer executes "
    "K-Means clustering (k=4) and Logistic Regression classification (80/20 train-test split, max_iter=1000) "
    "to generate trust scores and deployment readiness predictions. The pipeline completed full execution "
    "in 104.4 seconds on a lightweight VPS instance with 1.9 GB RAM, processing approximately 7,000 records "
    "across all sources. Evaluation metrics include a Silhouette score of 0.5276 (indicating moderate cluster "
    "separation) and Logistic Regression accuracy of 55.44% with F1 score of 46.59%, precision of 40.34%, "
    "and recall of 55.44%. The results demonstrate that a lightweight, Spark-based data pipeline can effectively "
    "unify disparate AI governance signals into structured, actionable metrics for organisational leadership."
)
EN_KEYWORDS = "data engineering; Agentic AI; trust assessment; Apache Spark; Medallion architecture; star schema; K-Means clustering; Logistic Regression; NLP contextual mapping"

EN_INTRODUCTION_PARAS = [
    "The rapid adoption of autonomous AI agents in professional services—law, consulting, accounting—has created an urgent need for structured trust assessment frameworks. Unlike traditional software, Agentic AI systems operate with significant autonomy: they triage client emails, monitor infrastructure, support financial reviews, and coordinate dynamic workflows across organisational boundaries. Without systematic monitoring and governance mechanisms, firms risk deploying agents whose outputs are unreliable, inconsistent, or misaligned with established standards.",
    "Current approaches to evaluating AI trustworthiness tend to focus on model-level metrics such as accuracy, precision, and recall in isolation. However, these metrics alone do not capture the broader organisational context: data provenance, source reliability, deployment frequency, and alignment with governance policies. There is no single, integrated data pipeline that combines operational telemetry, benchmark performance, and external validation signals into a unified trust assessment framework.",
    "Furthermore, most enterprise data pipelines are built on expensive, resource-intensive platforms such as Azure Data Factory or AWS Glue. Small-to-medium organisations and academic researchers often lack access to these platforms, creating a barrier to developing and validating AI governance solutions. A lightweight, open-source alternative built on Apache Spark would democratise access to these capabilities.",
    "This paper addresses that gap by designing and implementing a complete data engineering pipeline built on the Medallion architecture (Bronze-Silver-Gold) using Apache Spark. The pipeline ingests four heterogeneous data sources, applies contextual NLP mapping in the Silver layer, and executes clustering and classification models in the Gold layer to produce structured trust indicators. The final output follows a Star Schema analytical model suitable for downstream reporting and dashboarding. The research question is: How can a Medallion-architecture data pipeline transform heterogeneous AI deployment data into structured trust indicators through clustering and classification? The goal is to demonstrate that a lightweight, Spark-based pipeline can produce reproducible trust metrics from four distinct source types."
]
EN_LIT_REVIEW_PARAS = [
    "A data pipeline automates the movement of data from source systems through transformation into storage for analysis. Traditionally, Data Warehouses offered strong transactional guarantees and schema enforcement, but proved expensive to scale and poorly suited to semi-structured or streaming data. Conversely, Data Lakes provided cheap raw storage for unstructured data but lacked ACID support, leading to data swamps.",
    "The Data Lakehouse emerged as a middle ground, resolving this trade-off by layering transactional capabilities on top of lake storage. Armbrust et al. (2021) demonstrated that Lakehouse systems built on Apache Spark and Delta Lake can reduce data duplication, lower infrastructure costs, and support both business intelligence and machine learning workloads within a single platform. Their analysis underpins the storage paradigm adopted in this research.",
    "Zaharia et al. (2016) established the foundations of Apache Spark as a unified engine for large-scale data processing. Spark’s in-memory computation model enables iterative algorithms—such as K-Means clustering and logistic regression—to execute significantly faster than disk-based alternatives like MapReduce. This capability is critical for pipelines that must process thousands of records within seconds on constrained hardware.",
    "In the context of AI governance, Sculley et al. (2015) highlighted the hidden technical debt in machine learning systems, emphasizing that data management, feature engineering, and monitoring are often the most challenging aspects of ML deployment. Their work underscores the importance of building robust data pipelines before deploying models. Furthermore, Priem et al. (2022) developed OpenAlex as a fully open index of scholarly works. OpenAlex provides structured JSON metadata for academic publications, making it suitable as an external baseline for evaluating the research context and citation landscape of AI deployment domains.",
    "By synthesizing these perspectives—scalable Lakehouse storage, high-performance in-memory processing with Spark, and external baseline validation via OpenAlex—this paper defines a unified data framework. The Medallion architecture organizes data processing into three layers: Bronze (raw ingestion), Silver (cleansed and enriched), and Gold (business-level aggregates and analytical models), ensuring data quality and traceability at each stage."
]
EN_METHODOLOGY_PARAS = [
    "The project follows an adapted version of the Cross-Industry Standard Process for Data Mining (CRISP-DM) methodology, modified to emphasize data engineering concerns such as ingestion reliability, layer-specific data quality, and automated testing. The primary goal was to build a data engineering pipeline capable of ingesting four heterogeneous AI deployment data sources, transforming them through a layered architecture, and producing structured trust assessment indicators.",
    "The data preparation phase is organized into the three Medallion layers. The Bronze layer implements raw ingestion with schema enforcement. CSV files (2,500 deployment records and 2,000 log records) are read with explicitly defined schemas. JSON API responses (1,500 benchmark records and 1,000 OpenAlex metadata records) are flattened and normalized. All Bronze outputs are written as Parquet files with append mode, preserving the original data for auditability.",
    "The Silver layer applies cleansing, enrichment, and contextual mapping. Null values are filled with domain-appropriate defaults, and duplicate records are dropped based on primary keys. Schema validation ensures all required fields are correctly typed. Crucially, the OpenAlex metadata undergoes rule-based NLP contextual mapping, where research tokens are classified into governance categories (e.g., 'fairness', 'transparency', 'robustness') using keyword matching against a predefined taxonomy.",
    "The Gold layer focuses on analytical modeling and output generation. Numeric features (Leadership_Trust_Score, Productivity_Improvement, Record_Count) are assembled into vectors and scaled using StandardScaler. K-Means clustering (k=4) partitions records into deployment profile clusters. Logistic Regression classifies records into binary deployment readiness categories (80/20 train-test split, max_iter=1000). The Star Schema is materialized with Fact_AI_Deployment joined to Dim_Organization, Dim_Agent, and Dim_Research_Context."
]
EN_RESULTS_PARAS = [
    "The pipeline execution produces several key output artifacts that serve as interfaces for downstream consumers. The Bronze layer preserves raw ingested data as Parquet files. The Silver layer outputs cleansed and enriched data with NLP contextual mapping, representing the analytical-ready dataset. The Gold layer materializes the Star Schema, where the Fact_AI_Deployment table contains all quantitative measures and foreign keys referencing Dim_Organization, Dim_Agent, and Dim_Research_Context.",
    "The automated test suite, consisting of four validation modules, confirmed pipeline correctness. The Schema Test verified that all output tables contain the expected columns with correct data types. The Null Test asserted that critical fields (primary keys, foreign keys, metric values) contain zero null values. The Row Count Test verified that output record counts fall within expected ranges, with approximately 6,800 records persisting into the Gold layer after deduplication. The Model Artifact Test confirmed that trained models were exported correctly and are reproducible.",
    "In terms of performance, the pipeline processed all 7,000 source records in 104.4 seconds on a single VPS instance configured with 1.9 GB RAM. The K-Means clustering achieved a Silhouette Score of 0.5276, indicating moderate cluster separation among the four deployment profiles. The Logistic Regression model produced an accuracy of 55.44%, with an F1 Score of 46.59%, Precision of 40.34%, and Recall of 55.44%. While these metrics reflect the limited feature set of the synthetic data, they successfully establish a baseline for evaluating AI deployment readiness."
]
EN_CONCLUSION_PARAS = [
    "This paper demonstrated a complete data engineering pipeline for Agentic AI trust assessment, integrating four heterogeneous sources into a Medallion architecture using Apache Spark. The results validate that a lightweight, open-source data pipeline can effectively unify disparate governance signals into structured, actionable metrics.",
    "Key contributions include a repeatable ETL framework combining operational logs, benchmarks, and academic citations; rule-based NLP context mapping for semi-structured research metadata; and a Star Schema analytical model with integrated machine learning outputs. The pipeline completed its full execution efficiently on resource-constrained hardware, proving that scalable data engineering is achievable without significant infrastructure investment.",
    "Future work will focus on integrating real-time streaming sources via Spark Structured Streaming, expanding the NLP depth by replacing rule-based mapping with pre-trained transformer models, and deploying the Star Schema output to enterprise business intelligence platforms for interactive dashboarding."
]
EN_ACK = "The authors thank BINUS University for supporting this research for providing the domain context."

REFERENCES = [
    "[1] M. Armbrust, A. Ghodsi, R. Xin, and M. Zaharia, 'Lakehouse: A new generation of open platforms that unify data warehousing and advanced analytics,' in Proc. CIDR, 2021.",
    "[2] M. Zaharia et al., 'Apache Spark: A unified engine for big data processing,' Commun. ACM, vol. 59, no. 11, pp. 56-65, 2016.",
    "[3] J. Priem, H. Piwowar, and R. Orr, 'OpenAlex: A fully open index of scholarly works, authors, venues, institutions, and concepts,' arXiv:2205.01833, 2022.",
    "[4] D. Sculley et al., 'Hidden technical debt in machine learning systems,' in Proc. NeurIPS, 2015, pp. 2503-2511.",
    "[5] V. Bolon-Canedo, N. Sanchez-Marono, and A. Alonso-Betanzos, 'A review of feature selection methods on synthetic data,' Knowl. Inf. Syst., vol. 34, no. 3, pp. 483-519, 2013.",
    "[6] R. Kimball and M. Ross, The Data Warehouse Toolkit, 3rd ed. Hoboken, NJ: Wiley, 2013.",
    "[7] W. H. Inmon, Building the Data Warehouse, 4th ed. Indianapolis, IN: Wiley, 2005.",
    "[8] S. Lloyd, 'Least squares quantization in PCM,' IEEE Trans. Inf. Theory, vol. 28, no. 2, pp. 129-137, 1982.",
    "[9] D. R. Cox, 'The regression analysis of binary sequences,' J. R. Stat. Soc. B, vol. 20, no. 2, pp. 215-242, 1958.",
    "[10] J. Dean and S. Ghemawat, 'MapReduce: Simplified data processing on large clusters,' Commun. ACM, vol. 51, no. 1, pp. 107-113, 2008.",
    "[11] A. Ng, Machine Learning Yearning, 2018. [Online]. Available: https://www.mlyearning.org/",
]

# ─── Indonesian ──────────────────────────────────────────────────

ID_TITLE = "Pipeline Data Engineering untuk Penilaian Kepercayaan Agen AI di Jasa Profesional"
ID_ABSTRACT = (
    "Adopsi cepat agen AI otonom di jasa profesional—hukum, konsultan, akuntansi—menciptakan kebutuhan mendesak akan kerangka penilaian kepercayaan terstruktur. Tidak seperti perangkat lunak tradisional, sistem Agentic AI beroperasi dengan otonomi signifikan: mereka memilah email klien, memantau infrastruktur, mendukung tinjauan keuangan, dan mengoordinasi alur kerja dinamis lintas batas organisasi. Tanpa pemantauan dan mekanisme tata kelola yang sistematis, perusahaan berisiko menerapkan agen yang outputnya tidak andal, tidak konsisten, atau tidak selaras dengan standar yang berlaku. Makalah ini menangani kesenjangan tersebut dengan membangun pipeline data engineering yang mengintegrasikan telemetri operasional, data benchmark performa, dan metrik sitasi akademis eksternal ke dalam satu model analitis tunggal. Pipeline mengikuti arsitektur Medallion (Bronze-Silver-Gold) berbasis Apache Spark, memproses empat sumber heterogen: file CSV terstruktur dari Kaggle (2.500 rekor), benchmark JSON API dari SeedBench (1.500 rekor), log eksekusi agen tak terstruktur (2.000 rekor), dan metadata sitasi OpenAlex eksternal melalui REST API (1.000 rekor). Layer Silver menerapkan pemetaan konteks NLP berbasis aturan untuk mengkaya token konteks riset menggunakan heuristik kata kunci, sementara layer Gold menjalankan K-Means clustering (k=4) dan klasifikasi Logistic Regression (split latih/uji 80/20, max_iter=1000) untuk menghasilkan skor kepercayaan dan prediksi kesiapan penyejajaran. Pipeline menyelesaikan eksekusi penuh dalam 104,4 detik pada instance VPS ringan dengan RAM 1,9 GB, memproses sekitar 7.000 rekor dari semua sumber. Metrik evaluasi meliputi Skor Silhouette sebesar 0,5276 (menunjukkan pemisahan kluster moderat) dan akurasi Logistic Regression sebesar 55,44% dengan Skor F1 sebesar 46,59%, Presisi 40,34%, dan Recall 55,44%. Hasil menunjukkan bahwa pipeline data ringan berbasis Spark dapat menyatukan sinyal tata kelola AI yang heterogen menjadi metrik terstruktur yang dapat ditindaklanjuti bagi pimpinan organisasi."
)
ID_KEYWORDS = "data engineering; Agentic AI; penilaian kepercayaan; Apache Spark; arsitektur Medallion; skema bintang; K-Means clustering; Logistic Regression; pemetaan konteks NLP"

ID_INTRODUCTION_PARAS = [
    "Adopsi cepat agen AI otonom di jasa profesional—hukum, konsultan, akuntansi—menciptakan kebutuhan mendesak akan kerangka penilaian kepercayaan terstruktur. Tidak seperti perangkat lunak tradisional, sistem Agentic AI beroperasi dengan otonomi signifikan: mereka memilah email klien, memantau infrastruktur, mendukung tinjauan keuangan, dan mengoordinasi alur kerja dinamis lintas batas organisasi. Tanpa pemantauan dan mekanisme tata kelola yang sistematis, perusahaan berisiko menerapkan agen yang outputnya tidak andal, tidak konsisten, atau tidak selaras dengan standar yang berlaku.",
    "Pendekatan saat ini untuk menilai kepercayaan AI cenderung berfokus pada metrik model secara terisolasi, seperti akurasi, presisi, dan recall. Namun, metrik-metrik ini saja tidak menangkap konteks organisasi yang lebih luas: asal data, keandalan sumber, frekuensi penyejajaran, dan keselarasan dengan kebijakan tata kelola. Belum ada pipeline data terpadu yang menggabungkan sinyal-sinyal tata kelola, benchmark performa, dan validasi eksternal ke dalam kerangka penilaian kepercayaan yang tunggal dan kohesif.",
    "Lebih jauh, sebagian besar pipeline data enterprise dibangun di atas platform yang mahal dan intensif sumber daya seperti Azure Data Factory atau AWS Glue. Organisasi menengah kecil dan peneliti akademis sering kali tidak memiliki akses ke platform-platform ini, menciptakan hambatan untuk mengembangkan dan memvalidasi solusi tata kelola AI. Alternatif ringan berbasis sumber terbuka yang dibangun di atas Apache Spark akan mendemokratisasi akses terhadap kemampuan tersebut.",
    "Makalah ini menangani kesenjangan tersebut dengan merancang dan mengimplementasikan pipeline data engineering lengkap berbasis arsitektur Medallion (Bronze-Silver-Gold) menggunakan Apache Spark. Pipeline menyerap empat sumber data heterogen, menerapkan pemetaan konteks NLP di layer Silver, serta menjalankan model clustering dan klasifikasi di layer Gold untuk menghasilkan indikator kepercayaan terstruktur. Output akhir mengikuti model analitis Star Schema yang siap untuk pelaporan dan dasbor downstream."
]

ID_LIT_REVIEW_PARAS = [
    "Pipeline data mengotomatiskan perpindahan data dari sistem sumber melalui transformasi ke dalam penyimpanan untuk analisis. Secara tradisional, Data Warehouse menawarkan jaminan transaksi yang kuat dan penerapan skema, namun terbukti mahal untuk diskalakan dan tidak cocok untuk data semi-terstruktur atau streaming. Sebaliknya, Data Lake menyediakan penyimpanan mentah yang murah untuk data tak terstruktur namun tidak mendukung ACID, sehingga sering kali berubah menjadi rawa data.",
    "Data Lakehouse muncul sebagai jalan tengah, menyelesaikan trade-off ini dengan menumpuk kemampuan transaksi di atas penyimpanan lake. Armbrust et al. (2021) mendemonstrasikan bahwa sistem Lakehouse yang dibangun di atas Apache Spark dan Delta Lake dapat mengurangi duplikasi data, menurunkan biaya infrastruktur, serta mendukung beban kerja business intelligence dan machine learning dalam satu platform tunggal. Analisis mereka menjadi dasar paradigma penyimapanan yang diadopsi dalam penelitian ini.",
    "Zaharia et al. (2016) meletakkan fondasi Apache Spark sebagai mesin terpadu untuk pemrosesan data berskala besar. Model komputasi in-memory Spark memungkinkan algoritme iteratif—seperti K-Means clustering dan logistic regression—dieksekusi secara signifikan lebih cepat daripada alternatif berbasis disk seperti MapReduce. Kemampuan ini kritis bagi pipeline yang harus memproses ribuan rekor dalam hitungan detik di atas perangkat yang terbatas.",
    "Dalam konteks tata kelola AI, Sculley et al. (2015) menyoroti hutang teknis tersembunyi dalam sistem machine learning, menekankan bahwa manajemen data, rekayasa fitur, dan pemantauan sering kali merupakan aspek paling menantang dari penerapan ML. Karya mereka menegaskan pentingnya membangun pipeline data yang kokoh sebelum menerapkan model. Selain itu, Priem et al. (2022) mengembangkan OpenAlex sebagai indeks terbuka karya ilmiah yang menyediakan metadata JSON terstruktur, menjadikannya baseline eksternal yang sesuai untuk menilai konteks riset dan lanskap sitasi domain penerapan AI."
]

ID_METHODOLOGY_PARAS = [
    "Proyek ini mengikuti versi adaptasi metodologi Cross-Industry Standard Process for Data Mining (CRISP-DM), yang dimodifikasi untuk menekankan kekhawatiran data engineering seperti keandalan penyerapan, kualitas data spesifik layer, dan pengujian otomatis. Tujuan utamanya adalah membangun pipeline data engineering yang mampu menyerap empat sumber data penyejajaran AI heterogen, mentransformasinya melalui arsitektur berlayer, dan menghasilkan indikator kepercayaan terstruktur.",
    "Persiapan data diorganisasi ke dalam tiga layer Medallion. Layer Bronze menerapkan penyerapan mentah dengan penerapan skema. File CSV (2.500 rekor penyejajaran dan 2.000 rekor log) dibaca dengan skema yang didefinisikan secara eksplisit. Respons JSON API (1.500 rekor benchmark dan 1.000 metadata OpenAlex) ditata dan dinormalisasi. Semua output Bronze ditulis sebagai file Parquet dengan mode append, mempertahankan data mentah asli untuk keperluan audit.",
    "Layer Silver menerapkan pembersihan, pengayaan, dan pemetaan kontekstual. Nilai null diisi dengan nilai default yang sesuai domain (0 untuk numerik, 'Unknown' untuk kategorikal). Rekor duplikat dihapus berdasarkan primary key. Validasi skema memastikan semua kolom yang diperlukan hadil dengan tipe data yang benar. Secara krusial, metadata OpenAlex menjalani pemetaan konteks NLP berbasis aturan, di mana token riset diklasifikasikan ke dalam kategori tata kelola (misalnya 'keadilan', 'transparansi', 'robustness') menggunakan pencocokan kata kunci terhadap taksonomi yang sudah ditentukan.",
    "Layer Gold berfokus pada pemodelan analitis dan generasi output. Fitur numerik (Leadership_Trust_Score, Productivity_Improvement, Record_Count) dirangkai menjadi vektor dan diskalakan menggunakan StandardScaler. K-Means clustering (k=4) mempartisi rekor ke dalam kluster profil penyejajaran. Logistic Regression mengklasifikasikan rekor menjadi kategori kesiapan penyejajaran biner (split latih/uji 80/20, max_iter=1000). Skema Star diwujudkan dengan Fact_AI_Deployment yang terhubung ke Dim_Organization, Dim_Agent, dan Dim_Research_Context."
]

ID_RESULTS_PARAS = [
    "Pipeline menghasilkan artefak output kunci yang berfungsi sebagai antarmuka bagi konsumen downstream. Layer Bronze mempertahankan data mentah sebagai file Parquet. Layer Silver menghasilkan data yang telah dibersihkan dan diperkaya dengan pemetaan konteks NLP, yang merepresentasikan dataset siap analitis. Layer Gold mewujudkan Skema Star, di mana tabel Fact_AI_Deployment berisi semua ukuran kuantitatif dan foreign key yang merujuk ke Dim_Organization, Dim_Agent, dan Dim_Research_Context.",
    "Suite pengujian otomatis, yang terdiri dari empat modul validasi, mengonfirmasi kebenaran pipeline. Uji Skema memverifikasi bahwa semua tabel output mengandung kolom yang diharapkan dengan tipe data yang benar. Uji Null memastikan bahwa kolom-kolom kritis (primary key, foreign key, nilai metrik) tidak mengandung nilai nol. Uji Jumlah Baris memverifikasi bahwa jumlah rekor output berada dalam rentang yang diharapkan, dengan sekitar 6.800 rekor yang bertahan ke layer Gold setelah deduplikasi. Uji Artefak Model mengonfirmasi bahwa model yang telah dilatih diekspor dengan benar dan dapat direproduksi.",
    "Dari segi performa, pipeline memproses seluruh 7.000 rekor sumber dalam 104,4 detik pada satu instance VPS dengan RAM 1,9 GB. K-Means clustering mencapai Skor Silhouette sebesar 0,5276, yang mengindikasikan pemisahan kluster yang moderat di antara empat profil penyejajaran. Model Logistic Regression menghasilkan akurasi 55,44% dengan Skor F1 sebesar 46,59%, Presisi 40,34%, dan Recall 55,44%. Meskipun metrik-metrik ini mencerminkan keterbatasan fitur pada data sintetis, metrik ini berhasil menetapkan baseline untuk menilai kesiapan penyejajaran AI."
]

ID_CONCLUSION_PARAS = [
    "Makalah ini mendemonstrasikan pipeline data engineering yang lengkap untuk penilaian kepercayaan Agentic AI, mengintegrasikan empat sumber heterogen ke dalam arsitektur Medallion menggunakan Apache Spark. Hasil memvalidasi bahwa pipeline ringan berbasis sumber terbuka dapat menyatukan sinyal tata kelola yang berbeda-beda menjadi metrik terstruktur yang dapat ditindaklanjuti.",
    "Kontribusi utama meliputi kerangka ETL yang berulang yang menggabungkan log operasional, benchmark, dan sitasi akademis; pemetaan konteks NLP berbasis aturan untuk metadata riset semi-terstruktur; serta model analitis Star Schema dengan output machine learning terintegrasi. Pipeline menyelesaikan eksekusi penuh secara efisien di atas perangkat keras yang terbatas sumber dayanya, membuktikan bahwa rekayasa data yang diskalakan dapat dicapai tanpa investasi infrastruktur yang signifikan.",
    "Pekerjaan masa depan akan berfokus pada integrasi sumber streaming waktu nyata melalui Spark Structured Streaming, perluasan kedalaman NLP dengan mengganti pemetaan berbasis aturan menggunakan model transformer yang telah dilatih sebelumnya, serta penerapan output Skema Star ke platform business intelligence enterprise untuk dasbor interaktif."
]
ID_ACK = "Penulis berterima kasih kepada BINUS University atas dukungan riset ini."

# ─── Main builder ─────────────────────────────────────────────────

def build_paper(doc, is_english=True):
    """Modify template in-place: replace placeholders, then append content."""
    title = PAPER_TITLE if is_english else ID_TITLE
    abstract_txt = EN_ABSTRACT if is_english else ID_ABSTRACT
    kw_txt = EN_KEYWORDS if is_english else ID_KEYWORDS
    intro = EN_INTRODUCTION_PARAS if is_english else ID_INTRODUCTION_PARAS
    lit_review = EN_LIT_REVIEW_PARAS if is_english else ID_LIT_REVIEW_PARAS
    method = EN_METHODOLOGY_PARAS if is_english else ID_METHODOLOGY_PARAS
    results = EN_RESULTS_PARAS if is_english else ID_RESULTS_PARAS
    conclusion = EN_CONCLUSION_PARAS if is_english else ID_CONCLUSION_PARAS
    ack = EN_ACK if is_english else ID_ACK
    
    # ── Phase 1: Replace template placeholders ──
    kw_start_idx = None
    affiliation_group = 0  # 0=first group, 1=second group (clear only)
    for p_idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        
        # Replace paper title
        if "Paper Title (use style: paper title)" in txt:
            p.clear()
            r = p.add_run(title)
            set_run(r, 24, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Replace Authors
        if "Authors Name/s per 1st Affiliation" in txt:
            p.clear()
            p.style = doc.styles['Author']
            r = p.add_run(AUTHORS)
            set_run(r, 11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Replace first Affiliation lines
        if "line 1 (of Affiliation): dept. name of organization" in txt:
            if affiliation_group == 0:
                p.clear()
                p.style = doc.styles['Affiliation']
                r = p.add_run(AFFILIATION)
                set_run(r, 10, italic=True)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.clear()
        
        elif "line 2: name of organization" in txt:
            if affiliation_group == 0:
                p.clear()
                p.style = doc.styles['Affiliation']
                r = p.add_run("BINUS University, Jakarta, Indonesia 11480")
                set_run(r, 10, italic=True)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.clear()
        
        elif "line 3: City, Country" in txt:
            if affiliation_group == 0:
                p.clear()
                p.style = doc.styles['Affiliation']
                r = p.add_run("")
                set_run(r, 10, italic=True)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.clear()
        
        elif "line 4: e-mail" in txt:
            if affiliation_group == 0:
                p.clear()
                p.style = doc.styles['Affiliation']
                r = p.add_run(EMAIL)
                set_run(r, 9)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.clear()
        
        # Replace second Author (remove)
        if "2nd Affiliation (Author)" in txt:
            p.clear()
            affiliation_group = 1
        
        # Replace Subtitle
        if "Subtitle as needed" in txt:
            p.clear()
        
        # Replace conference header lines
        if "UTM Computing Proceedings" in txt:
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if "Innovations in Computing Technology" in txt:
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Replace Abstract
        if "Abstract\u2014" in txt or txt.startswith("Abstract\u2014"):
            p.clear()
            p.style = doc.styles['Abstract']
            r_bold = p.add_run("Abstract\u2014" if is_english else "Abstrak\u2014")
            set_run(r_bold, 9, bold=True)
            r = p.add_run(abstract_txt)
            set_run(r, 9)
            fmt_p(p, after=6, before=0)
        
        # Replace Keywords
        if "key words" in p.style.name:
            kw_start_idx = p_idx
            p.clear()
            p.style = doc.styles['key words']
            r = p.add_run(("Keywords\u2014" if is_english else "Kata kunci\u2014") + kw_txt)
            set_run(r, 9, bold=True, italic=True)
            fmt_p(p, after=6, before=0)
    
    if kw_start_idx is None:
        raise RuntimeError("Could not find Keywords paragraph")
    
    # ── Phase 2: Remove all paragraphs AFTER Keywords ──
    body = doc.element.body
    p_elements = body.findall(qn('w:p'))
    
    # Find Keywords element by style name (not text)
    kw_elem_idx = None
    for i, p_elem in enumerate(p_elements):
        # Check pPr for style reference
        pPr = p_elem.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                val = pStyle.get(qn('w:val'))
                if val == 'keywords':  # Internal style id for 'key words'
                    kw_elem_idx = i
                    break
    
    # Fallback: find by text content (only if style lookup failed)
    if kw_elem_idx is None:
        for i, p_elem in enumerate(p_elements):
            texts = p_elem.findall('.//'+qn('w:t'))
            full_text = ''.join(t.text or '' for t in texts)
            if 'Kata kunci' in full_text or 'Keywords\u2014' in full_text:
                kw_elem_idx = i
                break
    
    if kw_elem_idx is None:
        # Last resort: find last Normal paragraph before Heading 1
        for i, p_elem in enumerate(p_elements):
            texts = p_elem.findall('.//'+qn('w:t'))
            full_text = ''.join(t.text or '' for t in texts)
            if '1. Introduction' in full_text:
                kw_elem_idx = i - 1  # Previous paragraph should be Keywords
                break
    
    # Remove all paragraphs after Keywords
    for p_elem in p_elements[kw_elem_idx+1:]:
        body.remove(p_elem)
    
    # Also remove any tables after Keywords
    tbl_elements = body.findall(qn('w:tbl'))
    for tbl in tbl_elements:
        body.remove(tbl)
    
    # ── Phase 3: Ensure we're in a 2-column section ──
    # Find or create a 2-column section for body text
    last_section = doc.sections[-1]
    sect_pr = last_section._sectPr
    
    existing_cols = sect_pr.findall(qn('w:cols'))
    existing_type = sect_pr.findall(qn('w:type'))
    
    is_two_col = False
    for c in existing_cols:
        if c.get(qn('w:num')) == '2':
            is_two_col = True
            break
    
    if not is_two_col:
        # Make last section continuous + 2 columns
        if not existing_type:
            stype = etree.SubElement(sect_pr, qn('w:type'))
            stype.set(qn('w:val'), 'continuous')
        for c in existing_cols:
            sect_pr.remove(c)
        cols_el = etree.SubElement(sect_pr, qn('w:cols'))
        cols_el.set(qn('w:num'), '2')
        cols_el.set(qn('w:space'), '360')
    
    # ── Phase 4: Append content ──
    # 1. INTRODUCTION
    add_heading1(doc, "1. Introduction", 10)
    for p_text in intro:
        add_body(doc, p_text, sz=10, after=3)
    
    # 2. LITERATURE REVIEW
    add_heading1(doc, "2. Literature Review", 10)
    for p_text in lit_review:
        add_body(doc, p_text, sz=10, after=3)
    
    # 3. METHODOLOGY
    add_heading1(doc, "3. Methodology", 10)
    for p_text in method:
        add_body(doc, p_text, sz=10, after=3)
    
    # Diagrams: ETL + Use Case
    add_img(doc, ETL_PATH, 3.1)
    add_fig_caption(doc, "Figure 1: ETL Medallion Architecture Data Flow" if is_english else "Gambar 1: Aliran Data Arsitektur Medallion ETL")
    add_img(doc, UC_PATH, 3.1)
    add_fig_caption(doc, "Figure 2: Use Case Diagram" if is_english else "Gambar 2: Diagram Use Case")
    
    # 4. RESULTS & DISCUSSION
    add_heading1(doc, "4. Results and Discussion", 10)
    for p_text in results:
        add_body(doc, p_text, sz=10, after=3)
    
    # Star Schema ERD + Gantt
    add_img(doc, ERD_PATH, 3.1)
    add_fig_caption(doc, "Figure 3: Star Schema ERD Diagram" if is_english else "Gambar 3: Diagram ERD Star Schema")
    add_img(doc, GANTT_PATH, 3.1)
    add_fig_caption(doc, "Figure 4: Project Gantt Timeline" if is_english else "Gambar 4: Linimasa Proyek Gantt")
    
    # 5. CONCLUSION
    add_heading1(doc, "5. Conclusion", 10)
    for p_text in conclusion:
        add_body(doc, p_text, sz=10, after=3)
    
    # Acknowledgment
    add_heading1(doc, "Acknowledgment", 10)
    add_body(doc, ack, sz=9, indent=0, after=3)
    
    # References
    add_heading1(doc, "References", 10)
    for ref in REFERENCES:
        p = doc.add_paragraph(style='references')
        r = p.add_run(ref)
        set_run(r, 8)
        fmt_p(p, after=1, before=0, indent=0)

def main():
    for lang, out_path in [('EN', OUT_EN), ('ID', OUT_ID)]:
        shutil.copy2(TEMPLATE, out_path)
        doc = Document(out_path)
        build_paper(doc, is_english=(lang == 'EN'))
        doc.save(out_path)
        print(f"{lang} -> {out_path}")

if __name__ == '__main__':
    main()
