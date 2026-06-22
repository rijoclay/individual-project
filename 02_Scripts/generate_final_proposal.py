# -*- coding: utf-8 -*-
"""
generate_final_proposal.py — Proposal generator 10/10 Humanized.
Fixes:
  2) Penulisan human 10/10 (Abstract & Introduction = exact user text)
  3) Caption font size = 9pt (figures ONLY — tables unchanged)
  1) ERD baru sudah di generate_erd.py
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REPORT_DIR = os.path.abspath(os.path.join(BASE_DIR, '..', '07_Reports'))
VIZ_DIR   = os.path.abspath(os.path.join(BASE_DIR, '..', '06_Visualization'))
TEMPLATE  = r"C:\Users\richa\Documents\Sem.6\DatEng_Individual_Project\Materials\PROPOSAL_individual project.docx"
UC_PATH   = os.path.join(VIZ_DIR, 'use_case_diagram.png')
ERD_PATH  = os.path.join(VIZ_DIR, 'star_schema_erd.png')
GANTT_PATH= os.path.join(VIZ_DIR, 'gantt_chart.png')
ETL_PATH  = os.path.join(VIZ_DIR, 'etl_architecture.png')

BLACK = RGBColor(0,0,0)
TNR   = "Times New Roman"

# ─── helpers ────────────────────────────────────────────────────────

def _fmt(p, a=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, ls=1.5, indent=0.5):
    p.alignment = a
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = ls
    if indent > 0:
        p.paragraph_format.first_line_indent = Inches(indent)

def _fi(p, sz=12, b=False, i=False, c=BLACK):
    for r in p.runs:
        r.font.name = TNR; r.font.size = Pt(sz)
        r.bold = b; r.italic = i; r.font.color.rgb = c

def h1(d, t):
    p = d.add_paragraph()
    p.style = d.styles['Heading 1']
    p.clear(); _fmt(p, a=0, after=12, ls=1.15, indent=0)
    r = p.add_run(t); r.font.name=TNR; r.font.size=Pt(14); r.bold=True; r.font.color.rgb=BLACK

def h2(d, t):
    p = d.add_paragraph()
    p.style = d.styles['Heading 2']
    p.clear(); _fmt(p, a=0, after=6, ls=1.15, indent=0)
    r = p.add_run(t); r.font.name=TNR; r.font.size=Pt(12); r.bold=True; r.font.color.rgb=BLACK

def para(d, t, indent=0.5):
    p = d.add_paragraph(); _fmt(p, indent=indent); p.add_run(t); _fi(p); return p

def para0(d, t): return para(d, t, indent=0)

def fig_cap(d, t):
    """Figure caption — 9pt italic centered ONLY for figures."""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(t); r.font.name=TNR; r.font.size=Pt(9); r.italic=True; r.font.color.rgb=BLACK

def add_img(d, pth, w=5.0):
    p = d.add_paragraph(); p.alignment = 1
    p.add_run().add_picture(pth, width=Inches(w))

def toc(d, is_en):
    d.add_page_break()
    p = d.add_paragraph(); p.alignment=1
    r = p.add_run("Table of Contents" if is_en else "Daftar Isi")
    r.font.name=TNR; r.font.size=Pt(16); r.bold=True; r.font.color.rgb=BLACK
    en = [
        ("Abstract",1), ("1. Introduction",1),
        ("   1.1 Business Domain & Problem Identification",2),
        ("   1.2 Objectives",2), ("   1.3 Risk Analysis & Mitigation Plan",2),
        ("2. Literature Review",1), ("3. Methodology",1),
        ("   3.1 Methodology Selection & Justification",2),
        ("   3.2 Data Source Identification",2),
        ("   3.3 Data Ingestion, Cleansing & Processing",2),
        ("   3.4 Data Storage & Justification",2),
        ("   3.5 AI Algorithm Integration",2),
        ("4. System Design & Architecture",1),
        ("   4.1 Use Case Diagram Description",2),
        ("   4.2 System Architecture (ETL & Storage Flow)",2),
        ("   4.3 Database & Multidimensional Schema (Star Schema)",2),
        ("   4.4 Main Interface & Analytical Dashboard Design",2),
        ("   4.5 Testing & Validation Strategy",2),
        ("5. Conclusion",1), ("Appendix A: Project Timeline",1),
        ("Appendix B: Generative AI Use Log",1), ("References",1),
    ]
    id_ = [
        ("Abstrak",1), ("1. Pendahuluan",1),
        ("   1.1 Domain Bisnis & Identifikasi Masalah",2),
        ("   1.2 Tujuan Proyek",2), ("   1.3 Analisis Risiko & Rencana Mitigasi",2),
        ("2. Tinjauan Pustaka",1), ("3. Metodologi",1),
        ("   3.1 Pemilihan Metodologi & Justifikasi",2),
        ("   3.2 Identifikasi Sumber Data",2),
        ("   3.3 Ingesti Data, Pembersihan & Pemrosesan",2),
        ("   3.4 Penyimpanan Data & Justifikasi",2),
        ("   3.5 Integrasi Algoritma Kecerdasan Buatan",2),
        ("4. Desain Sistem & Arsitektur",1),
        ("   4.1 Deskripsi Use Case Diagram",2),
        ("   4.2 Arsitektur Sistem (Aliran ETL & Penyimpanan)",2),
        ("   4.3 Desain Database & Skema Multidimensi",2),
        ("   4.4 Antarmuka Utama & Dashboard Analitis",2),
        ("   4.5 Strategi Pengujian & Validasi",2),
        ("5. Kesimpulan",1), ("Lampiran A: Timeline Proyek",1),
        ("Lampiran B: Log Penggunaan AI Generatif",1), ("Daftar Pustaka",1),
    ]
    for t, lv in (en if is_en else id_):
        pe = d.add_paragraph(); pe.alignment = 0
        pe.paragraph_format.space_after=Pt(2); pe.paragraph_format.space_before=Pt(2)
        pe.paragraph_format.line_spacing=1.2
        pe.paragraph_format.left_indent=Inches(0.3 if lv==2 else 0)
        rr = pe.add_run(t); rr.font.name=TNR
        rr.font.size=Pt(12) if lv==1 else Pt(11)
        rr.bold = (lv==1); rr.font.color.rgb=BLACK

def tbl(d, rows, hdr, col_w=None):
    t = d.add_table(rows=len(rows)+1, cols=len(hdr))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hdr):
        t.rows[0].cells[i].text = ""; _fi(t.rows[0].cells[i].paragraphs[0], b=True)
    for idx, vals in enumerate(rows):
        for ci, v in enumerate(vals):
            t.rows[idx+1].cells[ci].text = str(v)
            _fi(t.rows[idx+1].cells[ci].paragraphs[0])
    d.add_paragraph()

def ref(d, t):
    p = d.add_paragraph(); p.alignment = 0
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    r = p.add_run(t); r.font.name=TNR; r.font.size=Pt(11); r.font.color.rgb=BLACK


# ═════════════════════════════════════════════════════════════════════
#  ENGLISH CONTENT  (Human 10/10)
# ═════════════════════════════════════════════════════════════════════

def content_en(d):
    toc(d, True)

    # ── ABSTRACT (exact user text) ──
    d.add_page_break()
    h1(d, "Abstract")
    para(d,
        "Agentic AI systems are starting to handle decisions that used to be made only by humans, "
        "from routing customer requests to prioritising operational incidents. Yet many organisations "
        "still rely on scattered logs and gut feeling to decide whether these agents are actually "
        "helping. This project tackles that gap by building a data pipeline that turns raw telemetry "
        "into clear performance evidence for leadership. Using Apache Spark and the Medallion "
        "architecture pattern (Bronze, Silver, Gold), the pipeline combines four data sources: a CSV "
        "dataset with 5,500 organisational records, JSON-based industry benchmarks, system execution "
        "logs, and OpenAlex academic citation data. The Silver layer focuses on data quality, applying "
        "industry-specific median imputation, schema standardisation, feature engineering, and "
        "regex-based NLP to map research papers to relevant business sectors. In the Gold layer, two "
        "machine learning models are deployed: a K-Means clustering model that groups organisations "
        "into performance tiers (Silhouette score 0.5276), and a tuned multiclass Logistic Regression "
        "model that predicts productivity categories with an accuracy of 55.44% and F1 of 46.59%. To avoid data leakage, "
        "clustering features are kept separate from the classifier inputs. The final pipeline produces "
        "repeatable, data-backed performance profiles that help leadership teams move away from "
        "anecdotal impressions and toward measurable, transparent evaluations of their Agentic AI "
        "systems.")

    # ── 1. INTRODUCTION (exact user text) ──
    d.add_page_break()
    h1(d, "1. Introduction")
    para(d,
        "Over the last few years, Agentic AI has moved from research labs into everyday business "
        "operations. Instead of simply following predefined rules, these systems can plan tasks, call "
        "external tools and APIs, and adjust their behaviour based on feedback. Organisations now use "
        "them to triage customer emails, monitor infrastructure, support financial reviews, and "
        "coordinate internal workflows. On paper, this should reduce costs and speed up decisions. "
        "In practice, many leadership teams still struggle to answer a basic question: is our AI "
        "actually performing well?")
    para(d,
        "The problem is not a lack of data, but a lack of structure. Execution logs live in text "
        "files. Operational metrics are exported as CSVs from dashboards. Industry benchmarks are only "
        "available through APIs, and academic evidence around AI trustworthiness is buried in external "
        "datasets like OpenAlex. Each source uses its own schema and naming conventions. Without a "
        "unified pipeline that can clean, align, and connect these signals, decision-makers are left "
        "comparing numbers manually or relying on subjective impressions. This makes it difficult to "
        "track performance over time, spot anomalies early, or justify further investment in Agentic "
        "AI.")
    para(d,
        "This project responds to that situation by designing and implementing a Spark-based data "
        "pipeline specifically tailored to Agentic AI evaluation. The pipeline follows the Medallion "
        "architecture, with a Bronze layer that ingests raw data, a Silver layer that enforces data "
        "quality and integration, and a Gold layer that applies machine learning. The end goal is not "
        "only to build a technically sound pipeline, but to give leadership teams performance profiles "
        "they can read, question, and use in real decisions about how, where, and whether to scale "
        "their Agentic AI deployments.")

    h2(d, "1.1 Business Domain & Problem Identification")
    para(d,
        "This project focuses on the Professional Services sector, which includes consulting firms, "
        "legal practices, IT service providers, and financial advisory organisations. These businesses "
        "run on expertise. Much of their value comes from how quickly and accurately they can interpret "
        "information, respond to clients, and manage risk. Because of this, they are early adopters of "
        "technologies that promise to augment human judgement rather than replace it entirely. Agentic "
        "AI fits naturally into this picture. It is being used to review documents, screen contracts, "
        "prioritise client queries, suggest data-driven recommendations, and automate routine "
        "operational tasks.")
    para(d,
        "When these systems work well, they can free consultants and analysts from repetitive work and "
        "allow them to focus on higher-value activities. However, the same autonomy that makes Agentic "
        "AI attractive also makes it harder to oversee. Leaders need to know whether agents are "
        "completing tasks reliably, how their behaviour differs across industries and use cases, and "
        "whether gains in productivity come at the cost of trust or stability. At the moment, there is "
        "no standard way for Professional Services firms to answer these questions. Telemetry is often "
        "stored in separate tools, logs, and exports, with no consistent schema or shared view.")
    para(d,
        "The core problem addressed in this project is therefore not just collecting data, but "
        "building a coherent data engineering framework around it. The goal is to bring together "
        "organisational context, agent configuration, operational metrics, benchmark baselines, and "
        "research signals into a single pipeline. Once this data is integrated and cleaned, it becomes "
        "possible to create performance profiles that show, for example, how leadership trust "
        "correlates with task success rates or how different agent types behave under similar workloads. "
        "This, in turn, gives Professional Services organisations a concrete basis for evaluating and "
        "governing their Agentic AI deployments.")

    h2(d, "1.2 Objectives")
    para(d, "The primary objective is to build a Spark-based data pipeline that turns fragmented Agentic "
        "AI telemetry into structured, analysable performance profiles. More specifically, the project "
        "aims to:", indent=0)
    for o in [
        "1. Build a Medallion-structured pipeline (Bronze, Silver, Gold) that ingests four "
        "heterogeneous sources: a 5,500-record CSV dataset, JSON benchmarks, system logs, and "
        "OpenAlex citation data.",
        "2. Implement automated cleansing using Spark UDFs and window functions, including "
        "industry-specific median imputation, deduplication, type standardisation, derived feature "
        "engineering, and regex-based NLP mapping of academic papers to business sectors.",
        "3. Deploy and validate a K-Means clustering model and a tuned Logistic Regression "
        "classifier, ensuring strict feature isolation between the two to prevent data leakage, "
        "and export the best model as a serialised production-ready artifact.",
    ]:
        para(d, o)

    h2(d, "1.3 Project Risk Analysis & Mitigation Plan")
    para(d,
        "Any data engineering project carries execution risks. Identifying them early means "
        "building mitigations into the architecture rather than patching them later.")
    tbl(d, [
        ("Data Quality & Schema Drift", "Impact: High\nProbability: Medium",
         "Median-based imputation for missing numericals; automated PySpark unit tests to validate schema integrity before Gold layer load."),
        ("Compute Limitations", "Impact: High\nProbability: High",
         "Running on 1.9GB RAM. Mitigated by tuning Spark partition counts, avoiding Python UDFs to prevent serialization overhead, and writing in batch chunks."),
        ("Stakeholder Alignment", "Impact: Medium\nProbability: Low",
         "Validating unsupervised clustering results (Silhouette Score) against leadership trust metrics ensures ML outputs map to business KPIs."),
    ], ("Identified Risk", "Impact / Probability", "Mitigation Strategy"))

    # ── 2. LITERATURE REVIEW ──
    d.add_page_break()
    h1(d, "2. Literature Review")
    para(d,
        "A data pipeline automates the movement of data from source systems through transformation "
        "steps and into storage for analysis. Traditionally, organisations relied on relational Data "
        "Warehouses, which offer strong transactional guarantees but are expensive to scale and poorly "
        "suited to semi-structured logs. Data Lakes provide cheap raw storage but lack ACID support "
        "and degrade under direct queries. The Data Lakehouse emerged as a middle ground: Parquet-"
        "formatted storage on commodity infrastructure, queried by engines like Apache Spark for both "
        "SQL analytics and machine learning (Armbrust et al., 2021).")
    para(d,
        "Three studies underpin this work. Armbrust et al. (2021) showed that Lakehouse architecture "
        "reduces data duplication and keeps a single source of truth for both dashboards and ML "
        "workloads. Zaharia et al. (2016) proved that Spark's in-memory processing handles complex "
        "joins across massive datasets, which is exactly the requirement for integrating scattered "
        "Agentic AI logs. Recent AI governance research highlights the importance of combining system "
        "telemetry with external baselines, such as OpenAlex citation data, to measure trust in "
        "autonomous agents (Priem et al., 2022). By synthesizing these three perspectives\u2014scalable "
        "Lakehouse storage, high-performance in-memory processing, and external baseline validation\u2014"
        "this project defines a unified data framework that not only stores AI logs efficiently but also maps "
        "them against academic trust metrics in a single analytical pipeline.")

    # ── 3. METHODOLOGY ──
    d.add_page_break()
    h1(d, "3. Methodology")
    para(d,
        "This section explains how the pipeline was designed, built, and validated. A clear "
        "methodology ensures the project stays focused, reproducible, and aligned with its objectives. "
        "We cover the project framework, the four data sources, the ETL process, storage rationale, "
        "and the ML models in the Gold layer.")
    para(d,
        "The pipeline follows the Medallion Architecture pattern. Bronze stores raw data unchanged. "
        "Silver cleans and integrates it. Gold applies analytics and machine learning. This "
        "separation keeps the code modular, handles bad data incrementally, and is standard practice "
        "in modern data engineering. The overall project lifecycle is managed using CRISP-DM.")

    h2(d, "3.1 Methodology Selection & Justification")
    para(d,
        "CRISP-DM was chosen because it is the most widely adopted framework for data mining "
        "projects, offering a structured yet flexible lifecycle. The six phases map directly to this "
        "project: Business Understanding covers the problem statement in Section 1; Data Understanding "
        "and Data Preparation map to Bronze and Silver layer processing; Modeling and Evaluation "
        "correspond to Gold layer analytics. The iterative nature of CRISP-DM is particularly "
        "valuable when integrating diverse sources with different schemas and quality profiles\u2014the "
        "preparation and modeling stages naturally inform each other.")

    h2(d, "3.2 Data Source Identification")
    para(d,
        "Four data sources capture the full operational context of Agentic AI deployments. They "
        "represent the three main categories organisations typically deal with: internal structured "
        "records, external benchmarks, system-level telemetry, and academic research context.")
    tbl(d, [
        ("S1", "Agentic AI Leadership Dataset", "CSV",
         "5,500 organisational records with autonomy levels, task success rates, and leadership trust scores."),
        ("S2", "External Leadership Benchmarks", "JSON",
         "Industry baseline targets for task success and productivity thresholds."),
        ("S3", "Agent Execution Logs", "CSV",
         "Granular system logs: CPU, memory, timestamps, response times, error counts."),
        ("S4", "OpenAlex Academic Publications", "API JSON",
         "25 scholarly publication records with citation counts and title keywords."),
    ], ("ID", "Data Source", "Format", "Description"))
    para(d,
        "Source S1 is the primary input: 5,500 CSV records, each representing one organisation\u2019s "
        "experience with an Agentic AI deployment. The 15 columns span organisational context "
        "(Record_ID, Organization_Name, Industry, Organization_Size, AI_Maturity_Level), agent "
        "configuration (Agent_Type, Use_Case_Area, Agent_Autonomy_Level, Decision_Making_Type, "
        "Context_Awareness_Score), and performance metrics (Task_Success_Rate, "
        "Productivity_Improvement_Percent, Response_Time_Seconds, Complexity_Score, "
        "Leadership_Trust_Score). Together they provide the multi-dimensional view needed for both "
        "clustering and classification in the Gold layer.")

    h2(d, "3.3 Data Ingestion, Cleansing & Processing")
    para(d,
        "Ingestion starts at the Bronze layer. CSV files are read with header inference and schema "
        "enforcement. JSON files are parsed with multi-line support for nested structures. The OpenAlex "
        "API response is consumed as standard JSON. Each ingested dataset is stamped with an "
        "ingestion timestamp and source identifier before being written to partitioned Parquet files.")
    para(d,
        "The Silver layer then cleans the data systematically. Duplicates are removed based on "
        "Record_ID. Numeric columns are cast to appropriate Spark types. Text fields are lowercased "
        "for consistency. Missing values are imputed using the per-industry median\u2014this preserves "
        "industry-specific distributions rather than applying a global average that could bias results. "
        "Execution log data is joined on Record_ID to derive efficiency features like "
        "Memory_Per_Message_MB. Finally, OpenAlex paper titles are processed through a regex NLP "
        "classifier that maps each publication to a business sector, and average citation counts per "
        "sector are joined to the organisational records.")

    h2(d, "3.4 Data Storage & Justification")
    para(d,
        "All data is stored in Apache Parquet across three Medallion layers. Parquet\u2019s columnar "
        "layout enables predicate pushdown and column projection during Spark SQL queries, reducing "
        "I/O. Dictionary encoding compresses repetitive strings. Schema enforcement prevents silent "
        "corruption. These characteristics make Parquet the industry standard for Lakehouse "
        "implementations.")

    h2(d, "3.5 AI Algorithm Integration")
    para(d,
        "The Gold layer deploys two ML models from PySpark MLlib. The first is unsupervised "
        "K-Means clustering, configured with k=3 clusters and 20 max iterations, operating on "
        "three normalised features: Task_Success_Rate, Productivity_Improvement_Percent, and "
        "Leadership_Trust_Score. These are assembled via VectorAssembler and standardised via "
        "StandardScaler. The Silhouette score of 0.5276 confirms moderate cluster separation, "
        "producing three distinguishable performance tiers: High, Average, and Low Performer.")
    para(d,
        "The second model is a supervised multiclass Logistic Regression classifier. It predicts "
        "Productivity_Category using exclusively system-level features: Context_Awareness_Score, "
        "Response_Time_Seconds, Task_Complexity_Score, Memory_Per_Message_MB, and "
        "CPU_Utilization_Percent. These features intentionally exclude clustering metrics to "
        "prevent data leakage. The model is tuned using CrossValidator with 3-fold CV over a "
        "regParam/elasticNetParam grid. The best model achieves 55.44% accuracy, 46.59% F1, "
        "40.34% precision, and 55.44% recall. It is serialised and exported to the Gold output "
        "directory for production use.")

    # ── 4. SYSTEM DESIGN ──
    d.add_page_break()
    h1(d, "4. System Design & Architecture")
    para(d,
        "The pipeline\u2019s architecture must handle the full lifecycle of Agentic AI telemetry: from "
        "raw, scattered data sources to a structured multidimensional schema. Apache Spark was chosen "
        "as the core engine because its distributed in-memory processing eliminates the "
        "Out-Of-Memory issues that single-node tools like Pandas hit when joining large execution "
        "logs with external JSON API responses.")
    para(d,
        "This chapter covers five components: the Use Case diagram, the ETL data flow, the Star "
        "Schema database design, the analytical dashboard interface, and the testing strategy that "
        "guards against data leakage between K-Means clustering and Logistic Regression "
        "classification.")

    h2(d, "4.1 Use Case Diagram Description")
    para(d,
        "Two actors interact with the system. The Data Engineer initiates pipeline execution, "
        "monitors processing status across all three Medallion layers, inspects logs for errors, "
        "and triggers reprocessing when data quality issues arise. The Data Analyst works with "
        "Gold layer outputs: running business intelligence queries, examining clustering profiles, "
        "evaluating classification predictions, and building visualisation dashboards for leadership "
        "reporting. The ML models run automatically as part of the pipeline, retraining on each "
        "fresh data batch and updating the serialised artifact.")
    if os.path.exists(UC_PATH):
        add_img(d, UC_PATH, 5.0)
        fig_cap(d, "Figure 4.1.1: Use Case Diagram of the Lakehouse System")

    h2(d, "4.2 System Architecture (ETL & Storage Flow)")
    para(d,
        "The pipeline follows a modular Medallion design. Raw datasets are stored locally and "
        "processed sequentially by Spark. Bronze produces audited Parquet archives. Silver generates "
        "cleansed, feature-engineered outputs. Gold produces aggregated analytical marts alongside "
        "ML predictions and the serialised model artifact. Each layer is an independent Spark job "
        "with its own input validation, transformation logic, and output verification steps. Logging "
        "at each stage provides full operational visibility.")
    if os.path.exists(ETL_PATH):
        add_img(d, ETL_PATH, 5.5)
        fig_cap(d, "Figure 4.2.1: ETL Medallion Architecture Data Flow")

    h2(d, "4.3 Database & Multidimensional Schema (Star Schema)")
    para(d,
        "The Gold layer is designed as a Multidimensional Database using the Star Schema pattern. "
        "This approach organises data into a central fact table surrounded by dimension tables, "
        "enabling efficient drill-down and slice-and-dice analysis across business perspectives. "
        "The central fact table, Fact_AI_Deployment, references three dimension tables: Dim_Organization "
        "(organisational attributes), Dim_Agent (AI agent configuration), and Dim_Research_Context "
        "(NLP-mapped scientific citation metrics). This design lets analysts query performance across "
        "multiple axes\u2014comparing trust scores by industry, agent type, or citation density\u2014without "
        "complex runtime joins. The ERD below maps each of the four data sources to their respective "
        "dimension tables, showing how raw input data is restructured into a multidimensional query model.")
    if os.path.exists(ERD_PATH):
        add_img(d, ERD_PATH, 5.5)
        fig_cap(d, "Figure 4.3.1: Star Schema ERD Diagram \u2014 Data Source to Multidimensional Mapping")

    h2(d, "4.4 Main Interface & Analytical Dashboard Design")
    para(d,
        "Three dashboards form the analytical interface. The first plots Leadership Trust Score "
        "against Productivity Improvement across all industries. The second shows clustering profile "
        "distributions. The third presents classification accuracy, confusion matrix statistics, and "
        "per-class performance breakdowns. These are intended for implementation in Power BI or "
        "Tableau, connecting directly to Gold layer Parquet files.")

    h2(d, "4.5 Testing & Validation Strategy")
    para(d,
        "An automated test suite runs after each pipeline execution. It checks four things: schema "
        "conformance (all expected derived columns exist with correct types), null assertions "
        "(critical fields contain zero nulls after Silver imputation), record integrity (exactly "
        "5,500 records preserved from Bronze to Gold), and model artifact validation (serialised "
        "model file exists at the expected path). All tests run on PySpark DataFrames to avoid "
        "pandas memory issues.")

    # ── 5. CONCLUSION ──
    d.add_page_break()
    h1(d, "5. Conclusion")
    para(d,
        "This proposal describes a data pipeline built to measure Agentic AI performance in real "
        "organisations. The Medallion architecture on Apache Spark integrates four heterogeneous "
        "data sources into a unified, analysable dataset. The Silver layer\u2019s cleansing "
        "operations\u2014industry-specific imputation, deduplication, feature engineering, and NLP "
        "contextual mapping\u2014produce a high-quality dataset. The Gold layer\u2019s ML models deliver "
        "actionable performance profiles while keeping clustering and classification features "
        "strictly isolated. The pipeline is production-ready: automated testing, serialised "
        "artifacts, and comprehensive logging ensure reliability and reproducibility.")

    # ── APPENDIX A ──
    d.add_page_break()
    h1(d, "Appendix A: Project Timeline")
    para0(d, "The project runs from 20 May to 10 July 2026. The Gantt chart below shows the task schedule and dependencies.")
    if os.path.exists(GANTT_PATH):
        add_img(d, GANTT_PATH, 6.0)
        fig_cap(d, "Figure A.1: Project Gantt Chart")

    # ── APPENDIX B ──
    h1(d, "Appendix B: Generative AI Use Log")
    tbl(d, [
        ("19 June 2026", "Code Refactoring",
         "Used Claude to optimize PySpark joins. Prompt: 'I have a memory bottleneck in my Silver layer join. How can I optimize this in PySpark without Pandas?'"),
        ("20 June 2026", "Proposal Draft",
         "Used ChatGPT to structure the methodology. Prompt: 'Help me outline a methodology section for a CRISP-DM data engineering project based on this rubric.'"),
        ("20 June 2026", "System Design",
         "Used Claude for Star Schema design. Prompt: 'Suggest a Star Schema design where the fact table tracks AI agent executions and dimensions include organization, agent config, and research context.'"),
    ], ("Date", "Task", "GAI Use / Prompt"))

    # ── REFERENCES ──
    h1(d, "References")
    for r in [
        "Armbrust, M., Ghodsi, A., Xin, R., & Zaharia, M. (2021). Lakehouse: A new generation of open platforms that unify data warehousing and advanced analytics. Proceedings of the 11th Conference on Innovative Data Systems Research (CIDR).",
        "Zaharia, M., Xin, R. S., Wendell, P., Das, T., Armbrust, M., Dave, A., Meng, X., Rosen, J., Venkataraman, S., Franklin, M. J., Ghodsi, A., Gonzalez, J., Shenker, S., & Stoica, I. (2016). Apache Spark: A unified engine for big data processing. Communications of the ACM, 59(11), 56\u201365.",
        "Priem, J., Piwowar, H., & Orr, R. (2022). OpenAlex: A fully open index of scholarly works, authors, venues, institutions, and concepts. Proceedings of the 26th International Conference on Science, Technology and Innovation Indicators (STI).",
        "Platt, J. C. (1999). Probabilistic outputs for support vector machines and comparisons to regularized likelihood methods. Advances in Large Margin Classifiers, 10(3), 61\u201374.",
        "Rousseeuw, P. J. (1987). Silhouettes: A graphical aid to the interpretation and validation of cluster analysis. Journal of Computational and Applied Mathematics, 20, 53\u201365.",
        "Bolon-Canedo, V., Sanchez-Marono, N., & Alonso-Betanzos, A. (2013). A review of feature selection methods on synthetic data. Knowledge and Information Systems, 34(3), 483\u2013519.",
        "Zhang, Y., & Zhang, Z. (2020). Data Lakehouse: A survey paper. arXiv preprint arXiv:2010.00874.",
        "Apache Software Foundation. (2023). Apache Spark MLlib: Scalable machine learning library. https://spark.apache.org/mllib/",
        "Wickham, H., & Grolemund, G. (2017). R for data science: Import, tidy, transform, visualize, and model data. O'Reilly Media.",
        "Kitchin, R. (2014). Big data, new epistemologies and paradigm shifts. Big Data & Society, 1(1), 1\u201312.",
    ]:
        ref(d, r)


# ═════════════════════════════════════════════════════════════════════
#  INDONESIAN CONTENT
# ═════════════════════════════════════════════════════════════════════

def content_id(d):
    toc(d, False)

    d.add_page_break()
    h1(d, "Abstrak")
    para(d,
        "Sistem Agentic AI kini mulai mengambil keputusan yang sebelumnya hanya bisa dilakukan "
        "manusia\u2014mulai dari merutekan permintaan pelanggan hingga memprioritaskan insiden "
        "operasional. Namun banyak organisasi masih mengandalkan log yang berserakan dan perkiraan "
        "kasar untuk memutuskan apakah agen-agen ini benar-benar membantu. Proyek ini menutup "
        "celah tersebut dengan membangun pipeline data yang mengubah telemetri mentah menjadi bukti "
        "kinerja yang jelas bagi pimpinan. Dengan Apache Spark dan pola arsitektur Medallion "
        "(Bronze, Silver, Gold), pipeline menggabungkan empat sumber data: dataset CSV berisi 5.500 "
        "catatan organisasi, tolok ukur industri berbasis JSON, log eksekusi sistem, dan data "
        "kutipan akademik OpenAlex. Layer Silver berfokus pada kualitas data, menerapkan imputasi "
        "median per industri, standarisasi skema, rekayasa fitur, dan NLP berbasis regex untuk "
        "memetakan paper penelitian ke sektor bisnis yang relevan. Di layer Gold, dua model "
        "pembelajaran mesin dijalankan: model klasterisasi K-Means yang mengelompokkan organisasi "
        "ke dalam tingkatan kinerja (skor Silhouette 0,5276), dan model Regresi Logistik multikelas "
        "yang memprediksi kategori produktivitas dengan akurasi 55,44% dan F1 46,59%. Untuk menghindari "
        "kebocoran data, fitur klasterisasi dijaga terpisah dari input klasifikasi. Pipeline "
        "menghasilkan profil kinerja yang berulang dan berbasis data, membantu tim kepemimpinan "
        "bergerak dari kesan intuitif menuju evaluasi yang terukur dan transparan terhadap sistem "
        "Agentic AI mereka.")

    d.add_page_break()
    h1(d, "1. Pendahuluan")
    para(d,
        "Dalam beberapa tahun terakhir, Agentic AI telah berpindah dari laboratorium penelitian ke "
        "operasi bisnis sehari-hari. Alih-alih hanya mengikuti aturan yang sudah ditentukan, sistem "
        "ini bisa merencanakan tugas, memanggil tools dan API eksternal, serta menyesuaikan "
        "perilaku berdasarkan umpan balik. Organisasi kini menggunakannya untuk menyeleksi email "
        "pelanggan, memantau infrastruktur, mendukung review keuangan, dan mengkoordinasi alur "
        "kerja internal. Di atas kertas, ini seharusnya menekan biaya dan mempercepat keputusan. "
        "Kenyataannya, banyak tim kepemimpinan masih kesulitan menjawab pertanyaan sederhana: "
        "apakah AI kami benar-benar berkinerja baik?")
    para(d,
        "Masalahnya bukan kurangnya data, melainkan kurangnya struktur. Log eksekusi tersembunyi di "
        "file teks. Metrik operasional tereksport sebagai CSV dari dashboard. Tolok ukur industri "
        "hanya tersedia melalui API, dan bukti akademis seputar kepercayaan AI terkubur di dataset "
        "eksternal seperti OpenAlex. Setiap sumber menggunakan skema dan konvensi penamaannya "
        "sendiri. Tanpa pipeline terpadu yang bisa membersihkan, menyelaraskan, dan menghubungkan "
        "sinyal-sinyal ini, para pengambil keputusan terpaksa membandingkan angka secara manual "
        "atau mengandalkan persepsi subjektif. Akibatnya, sulit melacak kinerja dari waktu ke "
        "waktu, mendeteksi anomali lebih awal, atau membenarkan investasi lebih lanjut dalam "
        "Agentic AI.")
    para(d,
        "Proyek ini merespons situasi tersebut dengan merancang dan mengimplementasikan pipeline "
        "data berbasis Spark yang dirancang khusus untuk evaluasi Agentic AI. Pipeline mengikuti "
        "arsitektur Medallion: layer Bronze menyimpan data mentah, layer Silver memastikan "
        "kualitas data dan integrasi, dan layer Gold menerapkan pembelajaran mesin. Tujuan "
        "akhirnya bukan sekadar membangun pipeline yang secara teknis benar, tetapi memberikan "
        "tim kepemimpinan profil kinerja yang bisa dibaca, dipertanyakan, dan digunakan dalam "
        "keputusan nyata tentang bagaimana, di mana, dan apakah akan memperluas deployment "
        "Agentic AI mereka.")

    h2(d, "1.1 Domain Bisnis & Identifikasi Masalah")
    para(d,
        "Proyek ini berfokus pada sektor Professional Services, yang mencakup firma konsultasi, "
        "praktik hukum, penyedia layanan TI, dan organisasi advisory keuangan. Bisnis-bisnis ini "
        "bergantung pada keahlian. Sebagian besar nilainya datang dari seberapa cepat dan akurat "
        "mereka bisa menginterpretasikan informasi, merespons klien, dan mengelola risiko. Karena "
        "itu, mereka merupakan pengguna awal teknologi yang berjanji untuk memperkuat penilaian "
        "manusia alih-alih menggantikannya sepenuhnya. Agentic AI pas dengan gambaran ini. Sistem "
        "ini digunakan untuk mereview dokumen, menyaring kontrak, memprioritaskan pertanyaan klien, "
        "memberikan rekomendasi berbasis data, dan mengotomasi tugas operasional rutin.")
    para(d,
        "Ketika sistem ini bekerja dengan baik, mereka bisa membebaskan konsultan dan analis dari "
        "pekerjaan berulang dan memungkinkan mereka fokus pada aktivitas bernilai lebih tinggi. "
        "Namun, otonomi yang membuat Agentic AI menarik juga membuatnya lebih sulit diawasi. "
        "Pemimpin perlu mengetahui apakah agen menyelesaikan tugas secara andal, bagaimana "
        "perilaku mereka berbeda lintas industri dan use case, dan apakah peningkatan produktivitas "
        "datang dengan mengorbankan kepercayaan atau stabilitas. Saat ini, belum ada cara standar "
        "bagi firma Professional Services untuk menjawab pertanyaan-pertanyaan ini. Telemetri "
        "sering tersimpan di tools, log, dan export yang terpisah, tanpa skema yang konsisten atau "
        "tampilan yang terpadu.")
    para(d,
        "Masalah pokok yang ditangani proyek ini bukan sekadar mengumpulkan data, tetapi "
        "membangun kerangka rekayasa data yang koheren di sekitarnya. Tujuannya adalah "
        "menggabungkan konteks organisasi, konfigurasi agen, metrik operasional, baseline "
        "benchmark, dan sinyal riset ke dalam satu pipeline. Begitu data ini terintegrasi dan "
        "dibersihkan, menjadi mungkin untuk membuat profil kinerja yang menunjukkan, misalnya, "
        "bagaimana kepercayaan pimpinan berkorelasi dengan tingkat keberhasilan tugas atau "
        "bagaimana tipe agen yang berbeda berperilaku di bawah beban kerja yang serupa. Ini "
        "pada gilirannya memberikan organisasi Professional Services dasar konkret untuk "
        "mengevaluasi dan mengatur deployment Agentic AI mereka.")

    h2(d, "1.2 Tujuan Proyek")
    para(d, "Tujuan utamanya adalah membangun pipeline data berbasis Spark yang mengubah telemetri "
        "Agentic AI yang terfragmentasi menjadi profil kinerja yang terstruktur dan dapat dianalisis. "
        "Secara lebih spesifik, proyek ini bertujuan untuk:", indent=0)
    for o in [
        "1. Membangun pipeline berstruktur Medallion (Bronze, Silver, Gold) yang mengumpulkan empat "
        "sumber heterogen: dataset CSV berisi 5.500 catatan, JSON benchmark, log sistem, dan data "
        "OpenAlex.",
        "2. Mengimplementasikan pembersihan otomatis menggunakan Spark UDF dan fungsi window, "
        "termasuk imputasi median per industri, deduplikasi, standarisasi tipe, rekayasa fitur "
        "turunan, dan pemetaan NLP berbasis regex dari paper akademik ke sektor bisnis.",
        "3. Menerapkan dan memvalidasi model klasterisasi K-Means dan classifier Regresi Logistik "
        "yang di-tune, memastikan isolasi fitur yang ketat antara keduanya untuk mencegah "
        "kebocoran data, dan mengekspor model terbaik sebagai artefak serialisasi siap produksi.",
    ]:
        para(d, o)

    h2(d, "1.3 Analisis Risiko & Rencana Mitigasi")
    para(d,
        "Setiap proyek rekayasa data membawa risiko eksekusi. Mengidentifikasinya sejak dini "
        "berarti membangun mitigasi ke dalam arsitektur alih-alih menambahkannya di kemudian hari.")
    tbl(d, [
        ("Kualitas Data & Pergeseran Skema", "Dampak: Tinggi\nProbabilitas: Sedang",
         "Imputasi median untuk nilai numerik yang hilang; unit test PySpark otomatis untuk memvalidasi integritas skema sebelum load ke layer Gold."),
        ("Keterbatasan Komputasi", "Dampak: Tinggi\nProbabilitas: Tinggi",
         "Berjalan pada RAM 1,9GB. Dimitigasi dengan menyesuaikan jumlah partition Spark, menghindari Python UDF untuk mencegah overhead serialisasi, dan menulis dalam batch chunk."),
        ("Alikan Pemangku Kepentingan", "Dampak: Sedang\nProbabilitas: Rendah",
         "Memvalidasi hasil klasterisasi tanpa pengawasan (Skor Silhouette) terhadap metrik kepercayaan pimpinan memastikan output ML selaras dengan KPI bisnis."),
    ], ("Risiko", "Dampak / Probabilitas", "Strategi Mitigasi"))

    d.add_page_break()
    h1(d, "2. Tinjauan Pustaka")
    para(d,
        "Pipeline data mengotomatiskan perpindahan data dari sistem sumber melalui langkah-langkah "
        "transformasi ke penyimpanan untuk analisis. Secara tradisional, organisasi mengandalkan "
        "Data Warehouse relasional, yang menawarkan jaminan transaksional kuat tetapi mahal "
        "diskalakan dan kurang cocok untuk log semi-terstruktur. Data Lake menyediakan penyimpanan "
        "mentah yang murah tetapi tidak memiliki dukungan ACID dan menurun kinerjanya saat "
        "di-query langsung. Data Lakehouse muncul sebagai jalan tengah: penyimpanan Parquet di "
        "atas infrastruktur komoditas, di-query oleh mesin seperti Apache Spark untuk analitik SQL "
        "sekaligus pembelajaran mesin (Armbrust et al., 2021).")
    para(d,
        "Tiga studi mendasari karya ini. Armbrust et al. (2021) menunjukkan bahwa arsitektur "
        "Lakehouse mengurangi duplikasi data dan mempertahankan satu sumber kebenaran untuk dasbor "
        "sekaligus workload ML. Zaharia et al. (2016) membuktikan bahwa pemrosesan in-memory Spark "
        "menangani join kompleks di dataset besar\u2014persis yang dibutuhkan untuk mengintegrasikan "
        "log Agentic AI yang tersebar. Riset tata kelola AI terbaru menyoroti pentingnya "
        "menggabungkan telemetri sistem dengan baseline eksternal, seperti data kutipan OpenAlex, "
        "untuk mengukur kepercayaan pada agen otonom (Priem et al., 2022). Dengan mensintesis ketiga perspektif "
        "ini\u2014penyimpanan Lakehouse yang skalabel, pemrosesan in-memory berkinerja tinggi, dan validasi "
        "baseline eksternal\u2014proyek ini mendefinisikan kerangka data terpadu yang tidak hanya menyimpan log "
        "AI secara efisien, tetapi juga memetakannya terhadap metrik kepercayaan akademis dalam satu "
        "pipeline analitis tunggal.")

    d.add_page_break()
    h1(d, "3. Metodologi")
    para(d,
        "Bagian ini menjelaskan bagaimana pipeline dirancang, dibangun, dan divalidasi. "
        "Metodologi yang jelas memastikan proyek tetap fokus, bisa direproduksi, dan selaras "
        "dengan tujuan. Kami membahas framework proyek, empat sumber data, proses ETL, alasan "
        "penyimpanan, dan model ML di layer Gold.")
    para(d,
        "Pipeline mengikuti pola Arsitektur Medallion. Bronze menyimpan data mentah tanpa "
        "perubahan. Silver membersihkan dan mengintegrasikan. Gold menerapkan analitik dan "
        "pembelajaran mesin. Pemisahan ini menjaga kode tetap modular, menangani data buruk "
        "secara bertahap, dan sudah jadi standar praktik di rekayasa data modern. Siklus hidup "
        "proyek secara keseluruhan dikelola menggunakan CRISP-DM.")

    h2(d, "3.1 Pemilihan Metodologi & Justifikasi")
    para(d,
        "CRISP-DM dipilih karena merupakan framework yang paling banyak diadopsi untuk proyek "
        "data mining, menawarkan siklus hidup yang terstruktur namun fleksibel. Enam fase "
        "tersebut langsung terpeta ke proyek ini: Pemahaman Bisnis mencakup pernyataan masalah "
        "di Bagian 1; Pemahaman Data dan Persiapan Data terpeta ke pemrosesan layer Bronze dan "
        "Silver; Pemodelan dan Evaluasi terkait dengan analitik layer Gold. Sifat iteratif "
        "CRISP-DM sangat berharga saat mengintegrasikan berbagai sumber dengan skema dan profil "
        "kualitas berbeda\u2014tahap persiapan dan pemodelan secara alami saling memberi informasi.")

    h2(d, "3.2 Identifikasi Sumber Data")
    para(d,
        "Empat sumber data menangkap konteks operasional penuh dari deployment Agentic AI. "
        "Mereka mewakili tiga kategori utama yang biasa dihadapi organisasi: catatan "
        "operasional terstruktur internal, benchmark eksternal, telemetri tingkat sistem, "
        "dan konteks riset akademik.")
    tbl(d, [
        ("S1", "Dataset Kepemimpinan Agentic AI", "CSV",
         "5.500 catatan organisasi dengan tingkat otonomi, keberhasilan tugas, dan skor kepercayaan pimpinan."),
        ("S2", "Tolok Ukur Kepemimpinan Eksternal", "JSON",
         "Target baseline industri untuk keberhasilan tugas dan ambang produktivitas."),
        ("S3", "Log Eksekusi Agen", "CSV",
         "Log sistem granular: utilisasi CPU, konsumsi memori, timestamp eksekusi, waktu respons, dan jumlah error."),
        ("S4", "Publikasi Akademik OpenAlex", "API JSON",
         "25 catatan publikasi ilmiah dengan jumlah kutipan dan kata kunci judul."),
    ], ("ID", "Sumber Data", "Format", "Deskripsi"))
    para(d,
        "Sumber S1 merupakan input utama: 5.500 record CSV, masing-masing mewakili pengalaman "
        "satu organisasi dengan deployment Agentic AI. Lima belas kolom mencakup konteks "
        "organisasi (Record_ID, Organization_Name, Industry, Organization_Size, "
        "AI_Maturity_Level), konfigurasi agen (Agent_Type, Use_Case_Area, Agent_Autonomy_Level, "
        "Decision_Making_Type, Context_Awareness_Score), dan metrik kinerja (Task_Success_Rate, "
        "Productivity_Improvement_Percent, Response_Time_Seconds, Complexity_Score, "
        "Leadership_Trust_Score).")

    h2(d, "3.3 Ingesti Data, Pembersihan & Pemrosesan")
    para(d,
        "Ingesti dimulai di layer Bronze. File CSV dibaca dengan inferensi header dan penegakan "
        "skema. File JSON dipars dengan dukungan multi-baris untuk struktur bersarang. Respons "
        "API OpenAlex dikonsumsi sebagai JSON standar. Setiap dataset yang diingesti diberi stempel "
        "waktu ingesti dan pengenal sumber sebelum ditulis ke file Parquet terpartisi.")
    para(d,
        "Layer Silver kemudian membersihkan data secara sistematis. Duplikat dihapus berdasarkan "
        "Record_ID. Kolom numerik dikonversi ke tipe Spark yang sesuai. Teks diubah ke huruf "
        "kecil untuk konsistensi. Nilai kosong diimputasi menggunakan median per industri\u2014ini "
        "mempertahankan distribusi spesifik industri alih-alih menerapkan rata-rata global yang "
        "bisa mem-bias hasil. Data log eksekusi digabungkan pada Record_ID untuk menurunkan "
        "fitur efisiensi seperti Memory_Per_Message_MB. Terakhir, judul paper OpenAlex "
        "diproses melalui classifier NLP regex yang memetakan setiap publikasi ke sektor bisnis, "
        "dan rata-rata jumlah kutipan per sektor digabungkan ke catatan organisasi.")

    h2(d, "3.4 Penyimpanan Data & Justifikasi")
    para(d,
        "Semua data disimpan dalam Apache Parquet di tiga layer Medallion. Tata letak kolom "
        "Parquet memungkinkan predicate pushdown dan column projection selama kueri Spark SQL, "
        "mengurangi I/O. Encoding kamus mengompresi string berulang. Penegakan skema mencegah "
        "korupsi senyap. Karakteristik ini menjadikan Parquet standar industri untuk "
        "implementasi Lakehouse.")

    h2(d, "3.5 Integrasi Algoritma Kecerdasan Buatan")
    para(d,
        "Layer Gold menjalankan dua model ML dari PySpark MLlib. Pertama, klasterisasi "
        "K-Means tanpa pengawasan, dikonfigurasi dengan k=3 klaster dan maksimal 20 iterasi, "
        "beroperasi pada tiga fitur ternormalisasi: Task_Success_Rate, "
        "Productivity_Improvement_Percent, dan Leadership_Trust_Score. Fitur-fitur ini dirakit "
        "melalui VectorAssembler dan distandarisasi via StandardScaler. Skor Silhouette 0,5276 "
        "mengkonfirmasi pemisahan klaster sedang, menghasilkan tiga tingkatan kinerja: High, "
        "Average, dan Low Performer.")
    para(d,
        "Model kedua adalah classifier Regresi Logistik multikelas terawasi. Model ini "
        "memprediksi Productivity_Category menggunakan fitur tingkat sistem secara eksklusif: "
        "Context_Awareness_Score, Response_Time_Seconds, Task_Complexity_Score, "
        "Memory_Per_Message_MB, dan CPU_Utilization_Percent. Fitur-fitur ini secara sengaja "
        "mengecualikan metrik klasterisasi untuk mencegah kebocoran data. Model di-tune "
        "menggunakan CrossValidator dengan 3-fold CV di atas grid regParam/elasticNetParam. "
        "Model terbaik mencapai akurasi 55,44%, F1 46,59%, presisi 40,34%, dan recall 55,44%. "
        "Model diserialisasi dan diekspor ke direktori output Gold untuk penggunaan produksi.")

    d.add_page_break()
    h1(d, "4. Desain Sistem & Arsitektur")
    para(d,
        "Arsitektur pipeline harus menangani siklus hidup penuh telemetri Agentic AI: dari "
        "sumber data yang mentah dan berserakan ke skema multidimensi yang terstruktur. Apache "
        "Spark dipilih sebagai mesin inti karena pemrosesan in-memory terdistribusinya "
        "menghilangkan masalah Out-Of-Memory yang dihadapi tool single-node seperti Pandas "
        "saat menggabungkan log eksekusi besar dengan respons JSON API eksternal.")
    para(d,
        "Bab ini mencakup lima komponen: diagram Use Case, alur data ETL, desain database "
        "Star Schema, antarmuka dashboard analitis, dan strategi pengujian yang menjaga "
        "pencegahan kebocoran data antara klasterisasi K-Means dan klasifikasi Regresi Logistik.")

    h2(d, "4.1 Deskripsi Use Case Diagram")
    para(d,
        "Dua aktor berinteraksi dengan sistem. Data Engineer memulai eksekusi pipeline, "
        "memantau status pemrosesan di ketiga layer Medallion, memeriksa log untuk error, dan "
        "memicu pemrosesan ulang saat masalah kualitas data terdeteksi. Data Analyst bekerja "
        "dengan output layer Gold: menjalankan kueri intelijen bisnis, memeriksa profil "
        "klasterisasi, mengevaluasi prediksi klasifikasi, dan membangun dashboard visualisasi "
        "untuk pelaporan pimpinan. Model ML berjalan otomatis sebagai bagian dari orkestrasi "
        "pipeline, melatih ulang pada setiap batch data baru dan memperbarui artefak terserialisasi.")
    if os.path.exists(UC_PATH):
        add_img(d, UC_PATH, 5.0)
        fig_cap(d, "Gambar 4.1.1: Diagram Use Case Sistem Lakehouse")

    h2(d, "4.2 Arsitektur Sistem (Aliran ETL & Penyimpanan)")
    para(d,
        "Pipeline mengikuti desain Medallion modular. Dataset mentah disimpan secara lokal "
        "dan diproses secara berurutan oleh Spark. Bronze menghasilkan arsip Parquet yang "
        "diaudit. Silver menghasilkan output yang telah dibersihkan dan direkayasa fitur. "
        "Gold menghasilkan mart analitis yang diagregasi bersama prediksi ML dan artefak "
        "model terserialisasi. Setiap layer merupakan job Spark independen dengan validasi "
        "input, logika transformasi, dan langkah verifikasi outputnya masing-masing. "
        "Pencatatan log di setiap tahap memberikan visibilitas operasional penuh.")
    if os.path.exists(ETL_PATH):
        add_img(d, ETL_PATH, 5.5)
        fig_cap(d, "Gambar 4.2.1: Aliran Data Arsitektur Medallion ETL")

    h2(d, "4.3 Desain Database & Skema Multidimensi (Star Schema)")
    para(d,
        "Layer Gold dirancang sebagai Database Multidimensi yang menggunakan pola Star Schema. "
        "Pendekatan ini mengorganisasikan data ke dalam satu tabel fakta pusat yang dikelilingi oleh "
        "tabel-tabel dimensi, memungkinkan analisis drill-down dan slice-and-dice yang efisien di berbagai "
        "perspektif bisnis. Tabel fakta pusat, Fact_AI_Deployment, mereferensikan tiga tabel dimensi: "
        "Dim_Organization (atribut organisasi), Dim_Agent (konfigurasi agen AI), dan Dim_Research_Context "
        "(metrik kutipan ilmiah hasil pemetaan NLP). Desain ini memungkinkan analis "
        "mengkueri metrik kinerja di berbagai sumbu analitis\u2014membandingkan skor kepercayaan "
        "menurut industri, tipe agen, atau kepadatan kutipan\u2014tanpa join multi-tabel yang "
        "kompleks di waktu kueri. ERD di bawah ini memetakan masing-masing dari keempat sumber data input "
        "ke tabel dimensinya masing-masing, menunjukkan bagaimana data mentah direstrukturisasi menjadi "
        "model kueri multidimensi.")
    if os.path.exists(ERD_PATH):
        add_img(d, ERD_PATH, 5.5)
        fig_cap(d, "Gambar 4.3.1: Diagram ERD Star Schema \u2014 Pemetaan Sumber Data ke Model Multidimensi")

    h2(d, "4.4 Antarmuka Utama & Desain Dashboard Analitis")
    para(d,
        "Tiga dashboard membentuk antarmuka analitis. Dashboard pertama memplot Skor "
        "Kepercayaan Kepemimpinan terhadap Peningkatan Produktivitas di semua industri. "
        "Dashboard kedua menampilkan distribusi profil klasterisasi. Dashboard ketiga "
        "menyajikan akurasi klasifikasi, statistik confusion matrix, dan breakdown "
        "kinerja per kelas. Dashboard ini diimplementasikan di Power BI atau Tableau, "
        "terhubung langsung ke file Parquet layer Gold.")

    h2(d, "4.5 Strategi Pengujian & Validasi")
    para(d,
        "Suite pengujian otomatis berjalan setelah setiap eksekusi pipeline. Empat aspek "
        "diperiksa: kesesuaian skema (semua kolom turunan yang diharapkan ada dengan tipe "
        "benar), asersi null (bidang kritis berisi nol null setelah imputasi Silver), "
        "integritas record (tepat 5.500 record terjaga dari Bronze ke Gold), dan validasi "
        "artefak model (file model terserialisasi ada di jalur yang diharapkan). Semua "
        "pengujian berjalan di PySpark DataFrame untuk menghindari masalah memori pandas.")

    d.add_page_break()
    h1(d, "5. Kesimpulan")
    para(d,
        "Proposal ini mendeskripsikan pipeline data yang dibangun untuk mengukur kinerja "
        "Agentic AI di organisasi nyata. Arsitektur Medallion di atas Apache Spark "
        "mengintegrasikan empat sumber data heterogen ke dalam dataset terpadu yang dapat "
        "dianalisis. Operasi pembersihan layer Silver\u2014imputasi spesifik industri, "
        "deduplikasi, rekayasa fitur, dan pemetaan kontekstual NLP\u2014menghasilkan dataset "
        "berkualitas tinggi. Model ML layer Gold memberikan profil kinerja yang dapat "
        "ditindaklanjuti sambil menjaga fitur klasterisasi dan klasifikasi tetap terpisah "
        "secara ketat. Pipeline siap produksi: pengujian otomatis, artefak terserialisasi, "
        "dan pencatatan log komprehensif memastikan keandalan dan reproducibility.")

    d.add_page_break()
    h1(d, "Lampiran A: Timeline Proyek")
    para0(d, "Proyek berlangsung dari 20 Mei hingga 10 Juli 2026. Diagram Gantt di bawah menampilkan jadwal tugas dan dependensi.")
    if os.path.exists(GANTT_PATH):
        add_img(d, GANTT_PATH, 6.0)
        fig_cap(d, "Gambar A.1: Diagram Gantt Proyek")

    h1(d, "Lampiran B: Log Penggunaan AI Generatif")
    tbl(d, [
        ("19 June 2026", "Refactoring Kode",
         "Pakai Claude untuk optimasi join PySpark. Prompt: 'Saya ada bottleneck memori di proses join Silver layer. Gimana cara optimasinya di PySpark tanpa pakai Pandas?'"),
        ("20 June 2026", "Draf Proposal",
         "Pakai ChatGPT untuk menstrukturisasi metodologi. Prompt: 'Tolong buatkan outline bagian metodologi untuk proyek data engineering CRISP-DM berdasarkan rubrik ini.'"),
        ("20 June 2026", "Desain Sistem",
         "Pakai Claude untuk brainstorming Star Schema. Prompt: 'Beri saran desain Star Schema di mana tabel fakta melacak eksekusi agen AI dan dimensinya mencakup organisasi, config agen, dan konteks riset.'"),
    ], ("Tanggal", "Tugas", "Penggunaan GAI / Prompt"))

    h1(d, "Daftar Pustaka")
    for r in [
        "Armbrust, M., Ghodsi, A., Xin, R., & Zaharia, M. (2021). Lakehouse: A new generation of open platforms that unify data warehousing and advanced analytics. Proceedings of the 11th Conference on Innovative Data Systems Research (CIDR).",
        "Zaharia, M., Xin, R. S., Wendell, P., Das, T., Armbrust, M., Dave, A., Meng, X., Rosen, J., Venkataraman, S., Franklin, M. J., Ghodsi, A., Gonzalez, J., Shenker, S., & Stoica, I. (2016). Apache Spark: A unified engine for big data processing. Communications of the ACM, 59(11), 56\u201365.",
        "Priem, J., Piwowar, H., & Orr, R. (2022). OpenAlex: A fully open index of scholarly works, authors, venues, institutions, and concepts. Proceedings of the 26th International Conference on Science, Technology and Innovation Indicators (STI).",
        "Platt, J. C. (1999). Probabilistic outputs for support vector machines and comparisons to regularized likelihood methods. Advances in Large Margin Classifiers, 10(3), 61\u201374.",
        "Rousseeuw, P. J. (1987). Silhouettes: A graphical aid to the interpretation and validation of cluster analysis. Journal of Computational and Applied Mathematics, 20, 53\u201365.",
        "Bolon-Canedo, V., Sanchez-Marono, N., & Alonso-Betanzos, A. (2013). A review of feature selection methods on synthetic data. Knowledge and Information Systems, 34(3), 483\u2013519.",
        "Zhang, Y., & Zhang, Z. (2020). Data Lakehouse: A survey paper. arXiv preprint arXiv:2010.00874.",
        "Apache Software Foundation. (2023). Apache Spark MLlib: Scalable machine learning library. https://spark.apache.org/mllib/",
        "Wickham, H., & Grolemund, G. (2017). R for data science: Import, tidy, transform, visualize, and model data. O'Reilly Media.",
        "Kitchin, R. (2014). Big data, new epistemologies and paradigm shifts. Big Data & Society, 1(1), 1\u201312.",
    ]:
        ref(d, r)


# ═════════════════════════════════════════════════════════════════════
#  MAIN
# ═════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    doc_en = Document(TEMPLATE)
    content_en(doc_en)
    p_en = os.path.join(REPORT_DIR, "Richard_Clay_Proposal_EN.docx")
    doc_en.save(p_en); print(f"EN  -> {p_en}")

    doc_id = Document(TEMPLATE)
    content_id(doc_id)
    p_id = os.path.join(REPORT_DIR, "Richard_Clay_Proposal_ID.docx")
    doc_id.save(p_id); print(f"ID  -> {p_id}")
